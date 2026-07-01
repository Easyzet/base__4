import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog
import pandas as pd
import threading
import time
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
import numpy as np
import math
import re
import os
import sys
import json
import urllib.request
import urllib.error

# ==========================================================================
#  Настройки автообновления с GitHub
#  Достаточно менять __version__ и пушить обновлённый файл в ветку —
#  публиковать релизы не требуется.
# --------------------------------------------------------------------------
__version__ = "1.0.0"                 # текущая версия приложения (увеличивайте при каждом обновлении)
GITHUB_OWNER = "Easyzet"            # владелец репозитория на GitHub
GITHUB_REPO = "base__4"       # имя репозитория
GITHUB_BRANCH = "main"               # ветка, из которой берётся обновление (обычно main или master)
# Путь к файлу внутри репозитория. Если файл лежит в корне — оставьте имя файла.
# Если в подпапке — укажите, например, "src/base.py".
REMOTE_SCRIPT_PATH = os.path.basename(sys.argv[0]) if sys.argv and sys.argv[0] else "base.py"
UPDATE_TIMEOUT = 6                    # таймаут сетевых запросов, сек
# ==========================================================================


def _raw_github_url():
    """URL raw-версии файла в репозитории для указанной ветки."""
    return (f"https://raw.githubusercontent.com/{GITHUB_OWNER}/{GITHUB_REPO}/"
            f"{GITHUB_BRANCH}/{REMOTE_SCRIPT_PATH}")


def _parse_version(v):
    """'v1.2.3' / '1.2.3' -> (1, 2, 3). При неудаче возвращает None."""
    if not v:
        return None
    s = str(v).strip().lstrip("vV")
    parts = re.split(r"[.\-_]", s)
    nums = []
    for p in parts:
        m = re.match(r"^\d+", p)
        if not m:
            break
        nums.append(int(m.group()))
    return tuple(nums) if nums else None


def _version_is_newer(remote, local):
    """True, если remote строго новее local."""
    rv, lv = _parse_version(remote), _parse_version(local)
    if rv is None or lv is None:
        # запасной вариант — строковое сравнение
        return str(remote).strip() != str(local).strip() and str(remote).strip() > str(local).strip()
    # выравниваем длину кортежей
    n = max(len(rv), len(lv))
    rv += (0,) * (n - len(rv))
    lv += (0,) * (n - len(lv))
    return rv > lv


class ColumnConfigDialog:
    def __init__(self, parent, columns: List[str]):
        self.parent = parent
        self.columns = columns.copy()
        self.result = None
        
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Настройка столбцов")
        self.dialog.geometry("400x500")
        self.dialog.transient(parent)
        self.dialog.grab_set()
        
        self.setup_ui()
        self.center_window()
        
    def setup_ui(self):
        # Заголовок
        title_label = tk.Label(self.dialog, text="Настройка столбцов таблицы", 
                              font=("Arial", 12, "bold"))
        title_label.pack(pady=10)
        
        # Фрейм для списка столбцов
        list_frame = tk.Frame(self.dialog)
        list_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # Listbox с прокруткой
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.columns_listbox = tk.Listbox(list_frame, yscrollcommand=scrollbar.set, 
                                         selectmode=tk.EXTENDED)
        self.columns_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.columns_listbox.yview)
        
        # Заполняем список
        for col in self.columns:
            self.columns_listbox.insert(tk.END, col)
        
        # Кнопки управления
        button_frame = tk.Frame(self.dialog)
        button_frame.pack(pady=10)
        
        tk.Button(button_frame, text="↑ Вверх", command=self.move_up).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="↓ Вниз", command=self.move_down).pack(side=tk.LEFT, padx=5)
        tk.Button(button_frame, text="Удалить", command=self.delete_column).pack(side=tk.LEFT, padx=5)
        
        # Кнопки действий
        action_frame = tk.Frame(self.dialog)
        action_frame.pack(pady=20)
        
        tk.Button(action_frame, text="Принять изменения", 
                 command=self.accept_changes, bg="green", fg="white").pack(side=tk.LEFT, padx=10)
        tk.Button(action_frame, text="Отмена", 
                 command=self.cancel).pack(side=tk.LEFT, padx=10)
    
    def move_up(self):
        selection = list(self.columns_listbox.curselection())
        if selection and min(selection) > 0:
            # Для множественного выделения
            items = [self.columns_listbox.get(i) for i in selection]
            for i in sorted(selection, reverse=True):
                self.columns_listbox.delete(i)
            
            new_pos = min(selection) - 1
            for item in items:
                self.columns_listbox.insert(new_pos, item)
                new_pos += 1
            
            # Восстанавливаем выделение
            for i in range(min(selection)-1, min(selection)-1 + len(items)):
                self.columns_listbox.selection_set(i)
    
    def move_down(self):
        selection = list(self.columns_listbox.curselection())
        if selection and max(selection) < self.columns_listbox.size() - 1:
            items = [self.columns_listbox.get(i) for i in selection]
            for i in sorted(selection, reverse=True):
                self.columns_listbox.delete(i)
            
            new_pos = max(selection) + 1 - len(selection) + 1
            for item in items:
                self.columns_listbox.insert(new_pos, item)
                new_pos += 1
            
            for i in range(max(selection)+1 - len(selection)+1, max(selection)+1 - len(selection)+1 + len(items)):
                self.columns_listbox.selection_set(i)
    
    def delete_column(self):
        selection = list(self.columns_listbox.curselection())
        if selection:
            if messagebox.askyesno("Подтверждение", f"Удалить {len(selection)} выбранных столбцов?"):
                for i in sorted(selection, reverse=True):
                    self.columns_listbox.delete(i)
    
    def accept_changes(self):
        self.result = []
        for i in range(self.columns_listbox.size()):
            self.result.append(self.columns_listbox.get(i))
        self.dialog.destroy()
    
    def cancel(self):
        self.result = None
        self.dialog.destroy()
    
    def center_window(self):
        self.dialog.update_idletasks()
        x = (self.dialog.winfo_screenwidth() // 2) - (400 // 2)
        y = (self.dialog.winfo_screenheight() // 2) - (500 // 2)
        self.dialog.geometry(f"400x500+{x}+{y}")

class ExcelViewerApp:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Для моего солнышка")
        self.root.geometry("900x600")
        
        self.excel_data: Dict[str, pd.DataFrame] = {}
        self.current_sheet = None
        self.current_columns = []
        self.current_column_name = None
        self.original_dataframe = None
        self.current_displayed_df = None
        self.file_path = None
        self.selected_columns = []
        self.criteria_frame = None
        self.criteria_table = None
        self.sheet_buttons = {}  # Исправлено: инициализируем sheet_buttons
        self.filter_stack = []  # Стек для хранения последовательности фильтров
        # Редактируемая/фильтруемая таблица
        self.df_full = None            # рабочая копия текущего листа (с правками)
        self.col_filters = {}          # {столбец: set(допустимых строковых значений)}
        self._active_col = None        # последний выбранный столбец (для поиска/замены)
        self._col_clipboard = None     # (имя, Series) — скопированный столбец
        self._cell_editor = None       # активный Entry-редактор ячейки
        self._search_state = None      # состояние поиска (для «Найти далее»)
        self._filter_popup = None      # единственное открытое окно фильтра
        self._ctx_menu = None          # активное контекстное меню
        
        self.setup_ui()
        self.center_window()
    
    def setup_ui(self):
        # Главный фрейм
        main_frame = tk.Frame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # ==== Панель вкладок с инструментами ====
        toolbar = ttk.Notebook(main_frame)
        toolbar.pack(fill=tk.X, pady=(0, 8))

        def _make_tab(title):
            f = tk.Frame(toolbar)
            toolbar.add(f, text=title)
            return f

        # --- Вкладка «Файл» ---
        tab_file = _make_tab("Файл")
        self.load_button = tk.Button(tab_file, text="Загрузить Excel",
                                     command=self.load_excel_file, bg="blue", fg="white")
        self.load_button.pack(side=tk.LEFT, padx=4, pady=6)
        self.save_excel_button = tk.Button(tab_file, text="Сохранить Excel",
                                           command=self.save_excel_file, state=tk.DISABLED)
        self.save_excel_button.pack(side=tk.LEFT, padx=4, pady=6)
        self.save_as_sheet_button = tk.Button(tab_file, text="Сохранить как лист",
                                              command=self.save_current_table_as_sheet,
                                              state=tk.DISABLED, bg="green", fg="white")
        self.save_as_sheet_button.pack(side=tk.LEFT, padx=4, pady=6)
        self.clear_button = tk.Button(tab_file, text="Очистить (закрыть файл)",
                                      command=self.clear_data, bg="red", fg="white")
        self.clear_button.pack(side=tk.LEFT, padx=4, pady=6)

        # --- Вкладка «Таблица» ---
        tab_table = _make_tab("Таблица")
        self.reset_filter_button = tk.Button(tab_table, text="Сбросить все фильтры",
                                             command=self.reset_all_filters, state=tk.DISABLED)
        self.reset_filter_button.pack(side=tk.LEFT, padx=4, pady=6)
        tk.Button(tab_table, text="+ Столбец",
                  command=self.add_column_dialog).pack(side=tk.LEFT, padx=4, pady=6)
        tk.Button(tab_table, text="+ Строка",
                  command=self.add_row).pack(side=tk.LEFT, padx=4, pady=6)
        tk.Button(tab_table, text="Поиск",
                  command=self.search_dialog).pack(side=tk.LEFT, padx=4, pady=6)
        tk.Button(tab_table, text="Замена",
                  command=self.replace_dialog).pack(side=tk.LEFT, padx=4, pady=6)
        self.config_button = tk.Button(tab_table, text="Настроить таблицу",
                                       command=self.configure_table, state=tk.DISABLED)
        self.config_button.pack(side=tk.LEFT, padx=4, pady=6)

        # --- Вкладка «Функции» ---
        tab_func = _make_tab("Функции")
        self.add_criteria_button = tk.Button(tab_func, text="Добавить критерии",
                                             command=self.add_criteria, state=tk.NORMAL,
                                             bg="orange", fg="white")
        self.add_criteria_button.pack(side=tk.LEFT, padx=4, pady=6)

        # --- Вкладка «Обновления» ---
        tab_upd = _make_tab("Обновления")
        self.check_updates_button = tk.Button(tab_upd, text="Проверить обновления",
                                              command=self.check_updates_manual)
        self.check_updates_button.pack(side=tk.LEFT, padx=4, pady=6)
        tk.Label(tab_upd, text=f"Текущая версия: {__version__}",
                 fg="#555555").pack(side=tk.LEFT, padx=(12, 4), pady=6)

        # ==== Строка прогресса и статуса ====
        status_frame = tk.Frame(main_frame)
        status_frame.pack(fill=tk.X, pady=(0, 8))
        self.progress = ttk.Progressbar(status_frame, mode='indeterminate')
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))
        self.status_label = tk.Label(status_frame, text="Готов к загрузке файла")
        self.status_label.pack(side=tk.RIGHT)
        
        # Панель листов
        sheets_frame = tk.Frame(main_frame)
        sheets_frame.pack(fill=tk.X, pady=(0, 10))
        
        tk.Label(sheets_frame, text="Доступные листы:", font=("Arial", 10, "bold")).pack(anchor=tk.W)
        
        self.sheets_frame = tk.Frame(sheets_frame)
        self.sheets_frame.pack(fill=tk.X, pady=(5, 0))
        
        # Основная рабочая область
        work_frame = tk.Frame(main_frame)
        work_frame.pack(fill=tk.BOTH, expand=True)
        
        # Левая панель (таблица)
        left_frame = tk.Frame(work_frame)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        self.root.bind("<Button-1>", self.hide_context_menus)
        self.root.bind("<Escape>", self._dismiss_menu)

        # Нижняя строка состояния под таблицей: слева — последнее действие, справа — число строк
        bottom_bar = tk.Frame(left_frame)
        bottom_bar.pack(side=tk.BOTTOM, fill=tk.X, pady=(3, 0))
        self.last_action_label = tk.Label(bottom_bar, text="", anchor="w", fg="#333333")
        self.last_action_label.pack(side=tk.LEFT)
        self.rowcount_label = tk.Label(bottom_bar, text="Строк: 0", anchor="e", fg="#333333")
        self.rowcount_label.pack(side=tk.RIGHT)

        # Таблица с прокруткой
        table_container = tk.Frame(left_frame)
        table_container.pack(fill=tk.BOTH, expand=True)
        table_container.configure(width=600, height=480)
        table_container.pack_propagate(False)
        
        # Горизонтальная прокрутка
        h_scrollbar = ttk.Scrollbar(table_container, orient=tk.HORIZONTAL)
        h_scrollbar.pack(side=tk.BOTTOM, fill=tk.X)
        
        # Вертикальная прокрутка
        v_scrollbar = ttk.Scrollbar(table_container, orient=tk.VERTICAL)
        v_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        # Treeview для таблицы с настройкой стиля
        style = ttk.Style()
        style.configure("Treeview", rowheight=25)
        style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        
        self.table = ttk.Treeview(table_container, 
                                 xscrollcommand=h_scrollbar.set,
                                 yscrollcommand=v_scrollbar.set)
        self.table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.h_scrollbar = h_scrollbar
        self.v_scrollbar = v_scrollbar
        self._colhl = []  # рамки подсветки выбранного столбца

        def _xscroll(*a):
            h_scrollbar.set(*a)
            self._highlight_active_column()

        def _yscroll(*a):
            v_scrollbar.set(*a)
            self._highlight_active_column()

        self.table.configure(xscrollcommand=_xscroll, yscrollcommand=_yscroll)
        self.table.bind("<Configure>", lambda e: self._highlight_active_column())
        
        # Настройка цветов строк для полосатости
        self.table.tag_configure('oddrow', background='#f0f0f0')
        self.table.tag_configure('evenrow', background='white')
        
        h_scrollbar.config(command=self.table.xview)
        v_scrollbar.config(command=self.table.yview)
        
        # Привязки: ЛКМ по заголовку — фильтр; ПКМ по заголовку — меню столбца;
        # двойной клик по ячейке — правка.
        self.table.bind('<Button-1>', self.on_table_left)
        self.table.bind('<Button-3>', self.on_table_right)
        self.table.bind('<Double-1>', self.on_table_double)

        # Прокрутка колёсиком мыши, когда курсор находится над таблицей.
        # Работает на Windows/macOS (<MouseWheel>) и Linux (<Button-4>/<Button-5>).
        def _table_wheel(event):
            if getattr(event, "num", None) == 4:          # Linux — вверх
                self.table.yview_scroll(-1, "units")
            elif getattr(event, "num", None) == 5:        # Linux — вниз
                self.table.yview_scroll(1, "units")
            elif event.delta:                             # Windows/macOS
                step = -1 if event.delta > 0 else 1
                if abs(event.delta) >= 120:               # Windows — кратно 120
                    step = int(-event.delta / 120)
                self.table.yview_scroll(step, "units")
            return "break"

        # Активируем колесо только когда курсор над таблицей, чтобы не мешать
        # прокрутке остальных областей.
        def _bind_wheel(_=None):
            self.table.bind_all("<MouseWheel>", _table_wheel)
            self.table.bind_all("<Button-4>", _table_wheel)
            self.table.bind_all("<Button-5>", _table_wheel)

        def _unbind_wheel(_=None):
            self.table.unbind_all("<MouseWheel>")
            self.table.unbind_all("<Button-4>")
            self.table.unbind_all("<Button-5>")

        self.table.bind("<Enter>", _bind_wheel)
        self.table.bind("<Leave>", _unbind_wheel)
        
        # Правая панель (данные столбца)
        right_frame = tk.Frame(work_frame, width=200)
        right_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=(10, 0))
        right_frame.pack_propagate(False)
        
        tk.Label(right_frame, text="Данные столбца:", font=("Arial", 10, "bold")).pack(anchor=tk.W)
        
        # Listbox для данных столбца с возможностью выделения
        column_scrollbar = tk.Scrollbar(right_frame)
        column_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Создаем Treeview для данных столбца
        self.column_data_table = ttk.Treeview(right_frame, yscrollcommand=column_scrollbar.set,
                                            selectmode=tk.EXTENDED, show='tree headings')
        self.column_data_table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        column_scrollbar.config(command=self.column_data_table.yview)

        # Настраиваем столбцы для таблицы данных столбца
        self.column_data_table['columns'] = ('row_num', 'value')
        self.column_data_table.column('#0', width=0, stretch=False)  # Скрываем первый столбец
        self.column_data_table.column('row_num', width=60, minwidth=50, anchor='center')
        self.column_data_table.column('value', width=140, minwidth=100)

        self.column_data_table.heading('row_num', text='№')
        self.column_data_table.heading('value', text='Значение')

        # Настраиваем цвета для четности строк
        self.column_data_table.tag_configure('oddrow', background='#f0f0f0')
        self.column_data_table.tag_configure('evenrow', background='white')
        self.column_data_table.tag_configure('empty_row', background='#ffe6e6')  # Для пустых строк
        
        # Контекстное меню для копирования
        self.context_menu = tk.Menu(self.root, tearoff=0)
        self.context_menu.add_command(label="Копировать выделенное", command=self.copy_selected_data)
        self.context_menu.add_command(label="Отобразить в таблице", command=self.filter_table_by_selection)
        self.column_data_table.bind("<Button-3>", self.show_context_menu)

    def add_criteria(self):
        """Окно выбора признаков: можно строить НЕСКОЛЬКО таблиц подряд,
        они складываются вертикально, и все вместе выгружаются в Word."""
        if not self.current_sheet:
            messagebox.showwarning("Предупреждение", "Загрузите данные перед настройкой критериев")
            return

        df = self.current_displayed_df if self.current_displayed_df is not None else self.excel_data[self.current_sheet]

        criteria_window = tk.Toplevel(self.root)
        criteria_window.title("Выбор критериев")
        criteria_window.geometry("1180x860")

        # Общая галочка: порядок строк/таблиц — по выделению или по порядку в списке
        order_bar = tk.Frame(criteria_window)
        order_bar.pack(fill=tk.X, padx=10, pady=(8, 0))
        criteria_window.order_by_sel = tk.BooleanVar(value=False)
        tk.Checkbutton(order_bar,
                       text="Упорядочивать по порядку выделения "
                            "(строки в таблице и таблицы при объединении)",
                       variable=criteria_window.order_by_sel,
                       font=("Arial", 9, "bold")).pack(side=tk.LEFT)
        # Состояние порядка выделения
        criteria_window.quant_sel_order = []
        criteria_window.merge_sel_order = []

        # Ручной выбор знаков после запятой для процентов в частотах
        pct_bar = tk.Frame(criteria_window)
        pct_bar.pack(fill=tk.X, padx=10, pady=(2, 0))
        criteria_window.pct_manual = tk.BooleanVar(value=False)
        criteria_window.pct_spin = ttk.Spinbox(pct_bar, from_=0, to=4, width=4)
        criteria_window.pct_spin.set(0)
        criteria_window.pct_spin.configure(state="disabled")

        def _toggle_pct_spin():
            criteria_window.pct_spin.configure(
                state="normal" if criteria_window.pct_manual.get() else "disabled")

        tk.Checkbutton(pct_bar, text="Ручной выбор знаков после запятой для частот",
                       variable=criteria_window.pct_manual,
                       command=_toggle_pct_spin, font=("Arial", 9)).pack(side=tk.LEFT)
        criteria_window.pct_spin.pack(side=tk.LEFT, padx=(4, 0))
        tk.Label(pct_bar, text="знаков").pack(side=tk.LEFT, padx=(3, 0))

        # --- Верх: две пронумерованные таблицы выбора ---
        selection_frame = tk.Frame(criteria_window)
        selection_frame.pack(fill=tk.X, padx=10, pady=(4, 4))

        qual_frame = tk.LabelFrame(selection_frame, text="Качественные признаки (группы — столбцы)",
                                   font=("Arial", 10, "bold"))
        qual_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))
        qual_scroll = ttk.Scrollbar(qual_frame)
        qual_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.qual_table = ttk.Treeview(qual_frame, yscrollcommand=qual_scroll.set,
                                       selectmode="extended", columns=("num", "column"),
                                       show="headings", height=9)
        self.qual_table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        qual_scroll.config(command=self.qual_table.yview)
        self.qual_table.heading("num", text="№")
        self.qual_table.column("num", width=45, minwidth=40, anchor="center", stretch=False)
        self.qual_table.heading("column", text="Признак")
        self.qual_table.column("column", width=260, minwidth=120)

        quant_frame = tk.LabelFrame(selection_frame,
                                    text="Признаки-строки (числовые -> Me [Q1; Q3]; качественные -> подсчёт)",
                                    font=("Arial", 10, "bold"))
        quant_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))
        quant_scroll = ttk.Scrollbar(quant_frame)
        quant_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.quant_table = ttk.Treeview(quant_frame, yscrollcommand=quant_scroll.set,
                                        selectmode="extended", columns=("num", "column", "sel"),
                                        show="headings", height=9)
        self.quant_table.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        quant_scroll.config(command=self.quant_table.yview)
        self.quant_table.heading("num", text="№")
        self.quant_table.column("num", width=45, minwidth=40, anchor="center", stretch=False)
        self.quant_table.heading("column", text="Признак")
        self.quant_table.column("column", width=220, minwidth=120)
        self.quant_table.heading("sel", text="№ выб.")
        self.quant_table.column("sel", width=55, minwidth=45, anchor="center", stretch=False)

        for idx, col in enumerate(df.columns, start=1):
            self.qual_table.insert("", tk.END, values=(idx, col))
            self.quant_table.insert("", tk.END, values=(idx, col, ""))

        # Отслеживаем порядок выделения строк-признаков
        self.quant_table.bind("<<TreeviewSelect>>",
                              lambda e: self._on_quant_select(criteria_window))

        # --- Блок выбора 95% ДИ (Клоппера-Пирсона) ---
        ci_frame = tk.LabelFrame(criteria_window,
                                 text="95% ДИ (Клоппера-Пирсона) — добавляется под таблицу частот",
                                 font=("Arial", 10, "bold"))
        ci_frame.pack(fill=tk.X, padx=10, pady=4)

        tk.Label(ci_frame, text="Признак:").pack(side=tk.LEFT, padx=(8, 2))
        self.ci_feature_cb = ttk.Combobox(ci_frame, state="readonly", width=24,
                                           values=["(не считать)"] + list(df.columns))
        self.ci_feature_cb.current(0)
        self.ci_feature_cb.pack(side=tk.LEFT, padx=2)

        tk.Label(ci_frame, text="Категория:").pack(side=tk.LEFT, padx=(8, 2))
        self.ci_category_cb = ttk.Combobox(ci_frame, state="readonly", width=18, values=[])
        self.ci_category_cb.pack(side=tk.LEFT, padx=2)

        tk.Label(ci_frame, text="Группа:").pack(side=tk.LEFT, padx=(8, 2))
        self.ci_group_cb = ttk.Combobox(ci_frame, state="readonly", width=26,
                                         values=["(весь набор данных)"])
        self.ci_group_cb.current(0)
        self.ci_group_cb.pack(side=tk.LEFT, padx=2)

        tk.Button(ci_frame, text="Очистить ДИ",
                  command=lambda: self.clear_ci_fields(criteria_window)).pack(side=tk.LEFT, padx=(10, 4))

        criteria_window.ci_feature_cb = self.ci_feature_cb
        criteria_window.ci_category_cb = self.ci_category_cb
        criteria_window.ci_group_cb = self.ci_group_cb

        # При выборе признака — обновляем список категорий
        self.ci_feature_cb.bind("<<ComboboxSelected>>",
                                lambda e: self._refresh_ci_categories(df, criteria_window))
        # При изменении выбора качественных признаков — обновляем список групп
        self.qual_table.bind("<<TreeviewSelect>>",
                             lambda e: self._refresh_ci_groups(df, criteria_window))

        # --- Кнопки ---
        button_frame = tk.Frame(criteria_window)
        button_frame.pack(fill=tk.X, padx=10, pady=4)
        tk.Button(button_frame, text="+ Таблица медиан (Me [Q1; Q3])",
                  command=lambda: self.add_result_table(df, criteria_window, "quant"),
                  bg="green", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(button_frame, text="+ Таблица частот (n, %)",
                  command=lambda: self.add_result_table(df, criteria_window, "count"),
                  bg="#2e7d32", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(button_frame, text="+ Таблица M±SD",
                  command=lambda: self.add_result_table(df, criteria_window, "mean"),
                  bg="#00695c", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(button_frame, text="Объединить выбранные",
                  command=lambda: self.merge_selected(criteria_window),
                  bg="#6a1b9a", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(button_frame, text="Закрыть",
                  command=criteria_window.destroy).pack(side=tk.RIGHT, padx=4)

        # --- Низ: прокручиваемая область со СТЕКОМ таблиц ---
        results_outer = tk.LabelFrame(criteria_window, text="Результаты (таблицы добавляются вниз одна за другой)",
                                      font=("Arial", 10, "bold"))
        results_outer.pack(fill=tk.BOTH, expand=True, padx=10, pady=(4, 10))

        res_canvas = tk.Canvas(results_outer, highlightthickness=0)
        res_vbar = ttk.Scrollbar(results_outer, orient="vertical", command=res_canvas.yview)
        res_hbar = ttk.Scrollbar(results_outer, orient="horizontal", command=res_canvas.xview)
        res_canvas.configure(yscrollcommand=res_vbar.set, xscrollcommand=res_hbar.set)
        res_vbar.pack(side=tk.RIGHT, fill=tk.Y)
        res_hbar.pack(side=tk.BOTTOM, fill=tk.X)
        res_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        results_inner = tk.Frame(res_canvas)
        res_canvas.create_window((0, 0), window=results_inner, anchor="nw")
        results_inner.bind("<Configure>",
                           lambda e: res_canvas.configure(scrollregion=res_canvas.bbox("all")))

        def _res_wheel(e):
            res_canvas.yview_scroll(int(-1 * (e.delta / 120)) if e.delta else 0, "units")

        res_canvas.bind("<Enter>", lambda e: res_canvas.bind_all("<MouseWheel>", _res_wheel))
        res_canvas.bind("<Leave>", lambda e: res_canvas.unbind_all("<MouseWheel>"))

        # --- Нижняя панель под полем таблиц (слева снизу) ---
        bottom_actions = tk.Frame(criteria_window)
        bottom_actions.pack(fill=tk.X, padx=10, pady=(0, 10))
        tk.Button(bottom_actions, text="Очистить результаты",
                  command=lambda: self.clear_results(criteria_window),
                  bg="#b71c1c", fg="white").pack(side=tk.LEFT, padx=(0, 6))
        tk.Button(bottom_actions, text="Выгрузить всё в Word",
                  command=lambda: self.export_to_word(df, criteria_window),
                  bg="blue", fg="white").pack(side=tk.LEFT, padx=6)

        # Ссылки/состояние окна
        criteria_window.qual_table = self.qual_table
        criteria_window.quant_table = self.quant_table
        criteria_window.results_inner = results_inner
        criteria_window.result_blocks = []
        criteria_window.df = df

    def _selected_columns(self, tree):
        """Имена выбранных признаков (в порядке таблицы). Имя — ячейка 'column'."""
        cols = []
        for item in tree.get_children():
            if item in tree.selection():
                name = tree.set(item, "column")
                if name:
                    cols.append(name)
        return cols

    def _on_quant_select(self, window):
        """Отслеживает порядок выделения строк-признаков и проставляет № выбора."""
        tree = window.quant_table
        current = set(tree.selection())
        order = window.quant_sel_order
        order[:] = [i for i in order if i in current]          # убрать снятые
        for i in tree.get_children():                          # добавить новые
            if i in current and i not in order:
                order.append(i)
        for i in tree.get_children():                          # обновить № выбора
            tree.set(i, "sel", str(order.index(i) + 1) if i in current else "")

    def _ordered_quant_columns(self, window):
        """Имена выбранных строк-признаков: по выделению или по таблице (по общей галочке)."""
        tree = window.quant_table
        sel = set(tree.selection())
        if window.order_by_sel.get():
            ordered = [i for i in window.quant_sel_order if i in sel]
            for i in tree.get_children():
                if i in sel and i not in ordered:
                    ordered.append(i)
        else:
            ordered = [i for i in tree.get_children() if i in sel]
        return [tree.set(i, "column") for i in ordered if tree.set(i, "column")]

    def _refresh_ci_categories(self, df, window):
        """Обновляет список категорий для выбранного признака 95% ДИ."""
        feat = window.ci_feature_cb.get()
        if not feat or feat == "(не считать)" or feat not in df.columns:
            window.ci_category_cb["values"] = []
            window.ci_category_cb.set("")
            return
        cats = sorted({self._pretty_value(v) for v in df[feat].dropna().astype(str)
                       if str(v).strip() != ''})
        window.ci_category_cb["values"] = cats
        if cats:
            window.ci_category_cb.current(0)

    def _refresh_ci_groups(self, df, window):
        """Обновляет список групп для 95% ДИ по выбранным качественным признакам."""
        options = ["(весь набор данных)"]
        for col in self._selected_columns(window.qual_table):
            for v in sorted({self._pretty_value(x) for x in df[col].dropna().astype(str)
                             if str(x).strip() != ''}):
                options.append(f"{col} = {v}")
        cur = window.ci_group_cb.get()
        window.ci_group_cb["values"] = options
        if cur in options:
            window.ci_group_cb.set(cur)
        else:
            window.ci_group_cb.current(0)

    def clear_ci_fields(self, window):
        """Очищает выбранные поля для 95% ДИ."""
        window.ci_feature_cb.set("(не считать)")
        window.ci_category_cb["values"] = []
        window.ci_category_cb.set("")
        window.ci_group_cb.set("(весь набор данных)")

    def add_result_table(self, df, window, kind):
        """Считает таблицу (медианы или частоты) и добавляет её блоком в стек результатов."""
        qual_selected = self._selected_columns(window.qual_table)
        quant_selected = self._ordered_quant_columns(window)
        if not qual_selected or not quant_selected:
            messagebox.showwarning("Предупреждение",
                                   "Выберите хотя бы один качественный признак (группа) "
                                   "и хотя бы один признак-строку")
            return

        if kind == "mean":
            columns, rows = self._compute_mean_rows(df, qual_selected, quant_selected)
            title = "M±SD: " + ", ".join(quant_selected) + "  |  группы: " + ", ".join(qual_selected)
            group_defs = []
            for col in qual_selected:
                for v in sorted([self._pretty_value(x) for x in df[col].dropna().astype(str).unique()
                                 if str(x).strip() != '']):
                    group_defs.append({"col": col, "value": v, "label": f"{col} = {v}"})
            block = {"kind": "quant", "kind_label": "Среднее (M ± SD)", "no_p": True,
                     "value_header": "M ± SD", "title": title, "columns": columns, "rows": rows,
                     "quant_features": list(quant_selected), "group_defs": group_defs,
                     "small_features": set(), "group_counts": {}, "p_results": []}
            window.result_blocks.append(block)
            self._render_block(window, block)
            return

        if kind == "quant":
            columns, rows, small_features, group_counts = self._compute_quant_rows(
                df, qual_selected, quant_selected)
            title = "Медианы: " + ", ".join(quant_selected) + "  |  группы: " + ", ".join(qual_selected)
            # Определения групп-столбцов (в порядке как в columns)
            group_defs = []
            for col in qual_selected:
                for v in sorted([self._pretty_value(x) for x in df[col].dropna().astype(str).unique()
                                 if str(x).strip() != '']):
                    group_defs.append({"col": col, "value": v, "label": f"{col} = {v}"})
            block = {"kind": "quant", "kind_label": "Медианы",
                     "title": title, "columns": columns, "rows": rows,
                     "quant_features": list(quant_selected), "group_defs": group_defs,
                     "small_features": set(small_features), "group_counts": group_counts,
                     "p_results": []}
            window.result_blocks.append(block)
            self._render_block(window, block)
            # Предложить заполнить столбец p
            self._ask_and_fill_pvalues(df, window, block)
            return
        else:
            # Если в поле «Группа» выбрано конкретное значение — строим таблицу
            # только по этой группе (по её столбцам), а не по всем значениям.
            grp = window.ci_group_cb.get()
            only_group = None
            if grp and grp != "(весь набор данных)" and " = " in grp:
                gcol, gval = grp.split(" = ", 1)
                only_group = (gcol.strip(), gval.strip())

            pct_manual = bool(window.pct_manual.get())
            try:
                pct_decimals = int(window.pct_spin.get())
            except Exception:
                pct_decimals = 0
            model = self._compute_count_structure(df, qual_selected, quant_selected, only_group,
                                                  pct_manual, pct_decimals)
            if only_group:
                title = "Частоты: " + ", ".join(quant_selected) + "  |  группа: " + grp
            else:
                title = "Частоты: " + ", ".join(quant_selected) + "  |  группы: " + ", ".join(qual_selected)
            ci = None
            feat = window.ci_feature_cb.get()
            cat = window.ci_category_cb.get()
            if feat and feat != "(не считать)" and cat:
                ci = self._compute_ci(df, feat, cat, grp)
            block = {"kind": "count", "kind_label": "Частоты",
                     "title": title, "model": model, "ci": ci, "p_results": []}

        window.result_blocks.append(block)
        self._render_block(window, block)
        if kind == "count":
            self._ask_and_fill_count_pvalues(df, window, block)

    def _ask_and_fill_count_pvalues(self, df, window, block):
        """Спрашивает, заполнять ли столбец p (точный критерий Фишера), и открывает окно."""
        groups = block["model"]["groups"]
        if len(groups) < 2:
            return  # для одной группы критерий не считается
        if not messagebox.askyesno(
                "Столбец p",
                "Заполнить столбец p по точному критерию Фишера?"):
            return
        self._open_count_pvalue_dialog(window, block)

    def _open_count_pvalue_dialog(self, window, block):
        """Окно расчёта p для таблицы частот (точный критерий Фишера / Фишера-Фримена-Холтона)."""
        model = block["model"]
        groups = model["groups"]
        features = [f["name"] for f in model["features"]]

        dlg = tk.Toplevel(self.root)
        dlg.title("Расчёт p")
        dlg.geometry("560x560")

        btns = tk.Frame(dlg)
        btns.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=8)
        content = self._scrollable_area(dlg)

        # Критерий (галочка только на одном)
        crit_frame = tk.LabelFrame(content, text="По какому критерию считать p",
                                   font=("Arial", 10, "bold"))
        crit_frame.pack(fill=tk.X, padx=10, pady=(10, 5))
        crit_vars = {}
        criteria = ["Точный критерий Фишера (ТКФ2 / Фишера-Фримена-Холтона)",
                    "Критерий хи-квадрат"]

        def select_criterion(chosen):
            for name, var in crit_vars.items():
                var.set(name == chosen)

        for i, name in enumerate(criteria):
            v = tk.BooleanVar(value=(i == 0))
            crit_vars[name] = v
            tk.Checkbutton(crit_frame, text=name, variable=v,
                           command=lambda n=name: select_criterion(n),
                           anchor="w").pack(fill=tk.X, padx=6)

        # Признаки-строки (по умолчанию все)
        rows_frame = tk.LabelFrame(content, text="Признаки для сравнения",
                                   font=("Arial", 10, "bold"))
        rows_frame.pack(fill=tk.X, padx=10, pady=(5, 5))
        feat_vars = {}
        for f in features:
            var = tk.BooleanVar(value=True)
            feat_vars[f] = var
            tk.Checkbutton(rows_frame, text=f, variable=var, anchor="w").pack(fill=tk.X, padx=6)

        # Выбор столбцов; «+» добавляет поле (для хи-квадрат можно 3+)
        cols_frame = tk.LabelFrame(content, text="Столбцы (группы) для сравнения",
                                   font=("Arial", 10, "bold"))
        cols_frame.pack(fill=tk.X, padx=10, pady=5)
        labels = [g["label"] for g in groups]
        combo_holder = tk.Frame(cols_frame)
        combo_holder.pack(fill=tk.X)
        col_combos = []

        def add_combo(default_idx=None):
            row = tk.Frame(combo_holder)
            row.pack(fill=tk.X, pady=2)
            tk.Label(row, text=f"Столбец {len(col_combos) + 1}:", width=10, anchor="w").pack(side=tk.LEFT)
            cb = ttk.Combobox(row, state="readonly", width=32, values=labels)
            cb.pack(side=tk.LEFT, padx=4)
            if default_idx is not None and default_idx < len(labels):
                cb.current(default_idx)

            def remove():
                if len(col_combos) <= 2:
                    return
                row.destroy()
                col_combos.remove(cb)

            tk.Button(row, text="−", width=2, command=remove).pack(side=tk.LEFT)
            col_combos.append(cb)

        add_combo(0)
        add_combo(1)
        tk.Button(cols_frame, text="+  Добавить столбец",
                  command=lambda: add_combo(min(len(col_combos), len(labels) - 1))).pack(anchor="w", padx=4, pady=2)
        tk.Label(cols_frame, text="(Фишер — два столбца; хи-квадрат — два и больше через «+»)",
                 fg="#555555").pack(anchor="w", padx=4)

        four_var = tk.BooleanVar(value=False)
        tk.Checkbutton(content, text="Показывать p с 4 знаками после запятой",
                       variable=four_var).pack(anchor="w", padx=12, pady=(4, 0))
        always_var = tk.BooleanVar(value=bool(block.get("always_label_criterion", False)))
        tk.Checkbutton(content, text="Всегда подписывать какой критерий использовался",
                       variable=always_var).pack(anchor="w", padx=12, pady=(0, 8))

        def do_calc():
            chosen = [n for n, v in crit_vars.items() if v.get()]
            if not chosen:
                messagebox.showwarning("Критерий", "Выберите критерий для расчёта p")
                return
            criterion = chosen[0]
            chi = (criterion == "Критерий хи-квадрат")
            sel_feats = [f for f, v in feat_vars.items() if v.get()]
            if not sel_feats:
                messagebox.showwarning("Признаки", "Выберите хотя бы один признак")
                return
            picked = []
            for cb in col_combos:
                l = cb.get()
                if l and l in labels and labels.index(l) not in picked:
                    picked.append(labels.index(l))
            if chi:
                if len(picked) < 2:
                    messagebox.showwarning("Столбцы", "Выберите минимум два РАЗНЫХ столбца")
                    return
            else:
                if len(picked) != 2:
                    messagebox.showwarning("Столбцы",
                                           "Для точного критерия Фишера нужно ровно два РАЗНЫХ столбца")
                    return
            four = four_var.get()
            res_pv, res_cr = {}, {}
            for f in model["features"]:
                if f["name"] not in sel_feats:
                    continue
                table = [[row["cells"][gi]["cnt"] for gi in picked] for row in f["rows"]]
                if chi:
                    res_pv[f["name"]] = self._fmt_pvalue(self._chi2_test(table), four=four)
                    res_cr[f["name"]] = "Критерий хи-квадрат"
                else:
                    res_pv[f["name"]] = self._fmt_pvalue(self._fisher_exact_2col(table), four=four)
                    res_cr[f["name"]] = ("Точный двусторонний критерий Фишера" if len(table) == 2
                                         else "Критерий Фишера-Фримена-Холтона")
            block.setdefault("p_results", []).append(
                {"pvalues": res_pv, "p_criteria": res_cr,
                 "compared": tuple(groups[gi]["value"] for gi in picked)})
            block["always_label_criterion"] = bool(always_var.get())
            self._rerender_all(window)
            dlg.destroy()

        tk.Button(btns, text="Рассчитать p", command=do_calc,
                  bg="green", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(btns, text="Отмена", command=dlg.destroy).pack(side=tk.RIGHT, padx=4)

    def _block_group_defs(self, block):
        """Возвращает определения групп-столбцов блока (список {col,value,label})."""
        if block["kind"] == "quant":
            return list(block.get("group_defs", []))
        return list(block["model"]["groups"])

    def _block_n_groups(self, block):
        return len(self._block_group_defs(block))

    def _p_column_specs(self, block):
        """Список столбцов p для отрисовки/экспорта:
        [{'header':..., 'pvalues':{feat:str}, 'p_criteria':{feat:crit}}].
        Если расчётов ещё не было — один пустой столбец 'p'
        (с прочерками для строк N<=4)."""
        results = block.get("p_results", [])
        specs = []
        n_groups = self._block_n_groups(block)
        for res in results:
            comp = res.get("compared")
            if comp and n_groups > 2:
                vals = comp if isinstance(comp, (list, tuple)) else [comp]
                header = "p (" + "; ".join(str(v) for v in vals) + ")"
            else:
                header = "p"
            specs.append({"header": header,
                          "pvalues": res.get("pvalues", {}),
                          "p_criteria": res.get("p_criteria", {})})
        if not specs:
            dashes = {f: "-" for f in block.get("small_features", set())}
            specs = [{"header": "p", "pvalues": dashes, "p_criteria": {}}]
        return specs

    def _p_criteria_index(self, specs, always=False):
        """Карта критерий->номер и упорядоченный список критериев.
        Индексы проставляются, если критериев больше одного ИЛИ включён режим
        «всегда подписывать критерий» (always=True) и критерий хотя бы один."""
        used = []
        for sp in specs:
            for feat, c in sp["p_criteria"].items():
                pt = sp["pvalues"].get(feat, "")
                if pt and pt != "-" and c and c not in used:
                    used.append(c)
        if len(used) > 1 or (always and len(used) >= 1):
            return {c: i + 1 for i, c in enumerate(used)}, used
        return {}, used

    def _block_as_features(self, block):
        """Преобразует блок в список признаков для объединённой таблицы + p/критерии
        (берётся первый набор расчётов p)."""
        feats = []
        pvalues = {}
        p_criteria = {}
        # Таблица отклонений (M ± SD) помечена no_p — для неё p не считается и не
        # должно подставляться при объединении, даже если признак с таким же именем
        # есть в другой таблице (иначе p дублировалось бы напротив отклонений).
        no_p = bool(block.get("no_p"))
        results = block.get("p_results", [])
        if results and not no_p:
            pv = results[0].get("pvalues", {})
            pc = results[0].get("p_criteria", {})
        elif no_p:
            pv = {}
            pc = {}
        else:
            pv = {f: "-" for f in block.get("small_features", set())}
            pc = {}
        if block["kind"] == "quant":
            n_groups = len(block.get("group_defs", []))
            for r in block["rows"]:
                name = r[0]
                cells = []
                for gi in range(n_groups):
                    n_val = r[1 + 2 * gi] if (1 + 2 * gi) < len(r) else ""
                    val = r[2 + 2 * gi] if (2 + 2 * gi) < len(r) else ""
                    cells.append({"N": n_val, "np": val})
                feats.append({"name": name, "binary": True, "no_p": no_p,
                              "rows": [{"category": "", "cells": cells}]})
                if not no_p:
                    if name in pv:
                        pvalues[name] = pv[name]
                    if name in pc:
                        p_criteria[name] = pc[name]
        else:
            for f in block["model"]["features"]:
                if no_p:
                    f = dict(f)
                    f["no_p"] = True
                feats.append(f)
                if not no_p:
                    if f["name"] in pv:
                        pvalues[f["name"]] = pv[f["name"]]
                    if f["name"] in pc:
                        p_criteria[f["name"]] = pc[f["name"]]
        return feats, pvalues, p_criteria

    def _refresh_merge_labels(self, window):
        """Обновляет подписи галочек объединения, показывая № порядка выделения."""
        order = window.merge_sel_order
        for b in window.result_blocks:
            cb = b.get("_sel_cb")
            if cb is None:
                continue
            if b.get("_selected") and b in order:
                try:
                    cb.config(text=f"Выбрать для объединения ({order.index(b) + 1})")
                except Exception:
                    pass
            else:
                try:
                    cb.config(text="Выбрать для объединения")
                except Exception:
                    pass

    def merge_selected(self, window):
        """Объединяет выбранные таблицы в одну (если столбцы-группы совпадают).
        Порядок таблиц — по выделению или по списку (по общей галочке)."""
        if window.order_by_sel.get():
            selected = [b for b in window.merge_sel_order if b.get("_selected")]
            # на случай, если что-то выбрано мимо учёта порядка
            for b in window.result_blocks:
                if b.get("_selected") and b not in selected:
                    selected.append(b)
        else:
            selected = [b for b in window.result_blocks if b.get("_selected")]
        if len(selected) < 2:
            messagebox.showwarning("Объединение",
                                   "Отметьте минимум две таблицы галочкой «Выбрать для объединения»")
            return
        base_defs = self._block_group_defs(selected[0])
        base_labels = [g["label"] for g in base_defs]
        for b in selected[1:]:
            if [g["label"] for g in self._block_group_defs(b)] != base_labels:
                messagebox.showwarning(
                    "Объединение невозможно",
                    "У выбранных таблиц разные столбцы (группирующие признаки). "
                    "Объединение отменено.")
                return

        features = []
        pvalues = {}
        p_criteria = {}
        for b in selected:
            f, pv, pc = self._block_as_features(b)
            features += f
            pvalues.update(pv)
            p_criteria.update(pc)

        model = {"groups": base_defs, "features": features}
        title = "Объединённая таблица  |  группы: " + ", ".join(
            sorted({g["col"] for g in base_defs}))
        merged = {"kind": "count", "merged": True, "kind_label": "Объединённая таблица",
                  "title": title, "model": model, "ci": None,
                  "value_header": "Me [Q1; Q3]/n(%)",
                  "p_results": [{"pvalues": pvalues, "p_criteria": p_criteria,
                                 "compared": None}]}
        # снять отметки с исходных таблиц
        for b in selected:
            b["_selected"] = False
        window.merge_sel_order = []
        window.result_blocks.append(merged)
        self._rerender_all(window)

    def clear_results(self, window):
        """Удаляет все таблицы из стека результатов."""
        window.result_blocks = []
        window.merge_sel_order = []
        for child in list(window.results_inner.winfo_children()):
            child.destroy()

    def _rerender_all(self, window):
        """Перерисовывает все блоки заново (например, после заполнения p)."""
        for child in list(window.results_inner.winfo_children()):
            child.destroy()
        for block in window.result_blocks:
            self._render_block(window, block)
        self._refresh_merge_labels(window)

    def _scrollable_area(self, parent):
        """Создаёт прокручиваемую (вертикально) область внутри parent и возвращает
        внутренний фрейм, куда можно класть виджеты."""
        container = tk.Frame(parent)
        container.pack(side=tk.TOP, fill=tk.BOTH, expand=True)
        canvas = tk.Canvas(container, highlightthickness=0)
        vbar = ttk.Scrollbar(container, orient=tk.VERTICAL, command=canvas.yview)
        canvas.configure(yscrollcommand=vbar.set)
        vbar.pack(side=tk.RIGHT, fill=tk.Y)
        canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        inner = tk.Frame(canvas)
        win_id = canvas.create_window((0, 0), window=inner, anchor="nw")
        inner.bind("<Configure>",
                   lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.bind("<Configure>", lambda e: canvas.itemconfig(win_id, width=e.width))

        def _wheel(e):
            canvas.yview_scroll(int(-1 * (e.delta / 120)) if e.delta else 0, "units")

        canvas.bind("<Enter>", lambda e: canvas.bind_all("<MouseWheel>", _wheel))
        canvas.bind("<Leave>", lambda e: canvas.unbind_all("<MouseWheel>"))
        return inner

    def _ask_and_fill_pvalues(self, df, window, block):
        """Спрашивает, заполнять ли столбец p, и при согласии открывает окно выбора."""
        if not messagebox.askyesno(
                "Столбец p",
                "Заполнить столбец p по критерию Манна-Уитни?"):
            return
        self._open_pvalue_dialog(df, window, block)

    def _open_pvalue_dialog(self, df, window, block):
        """Окно выбора признаков (строк) и двух столбцов для расчёта p (Манна-Уитни)."""
        features = block.get("quant_features", [])
        groups = block.get("group_defs", [])
        if len(groups) < 2:
            messagebox.showwarning("Недостаточно групп",
                                   "Для критерия Манна-Уитни нужно минимум два столбца (группы).")
            return

        dlg = tk.Toplevel(self.root)
        dlg.title("Расчёт p")
        dlg.geometry("560x560")

        # Кнопки закреплены снизу, остальное — в прокручиваемой области
        btns = tk.Frame(dlg)
        btns.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=8)
        content = self._scrollable_area(dlg)

        # --- Критерий для расчёта p (галочка только на одном) ---
        crit_frame = tk.LabelFrame(content, text="По какому критерию считать p",
                                   font=("Arial", 10, "bold"))
        crit_frame.pack(fill=tk.X, padx=10, pady=(10, 5))
        crit_vars = {}
        criteria = ["Критерий Манна-Уитни", "Критерий Краскелла-Уоллиса"]

        def select_criterion(chosen):
            for name, var in crit_vars.items():
                var.set(name == chosen)

        for i, name in enumerate(criteria):
            v = tk.BooleanVar(value=(i == 0))
            crit_vars[name] = v
            tk.Checkbutton(crit_frame, text=name, variable=v,
                           command=lambda n=name: select_criterion(n),
                           anchor="w").pack(fill=tk.X, padx=6)

        # --- Признаки-строки (по умолчанию все выбраны) ---
        rows_frame = tk.LabelFrame(content, text="Признаки (строки) для сравнения",
                                   font=("Arial", 10, "bold"))
        rows_frame.pack(fill=tk.X, padx=10, pady=(5, 5))
        feat_vars = {}
        for f in features:
            var = tk.BooleanVar(value=True)
            feat_vars[f] = var
            tk.Checkbutton(rows_frame, text=f, variable=var, anchor="w").pack(fill=tk.X, padx=6)

        # --- Выбор столбцов (групп); «+» добавляет поле (для Краскелла-Уоллиса) ---
        cols_frame = tk.LabelFrame(content, text="Столбцы (группы) для сравнения",
                                   font=("Arial", 10, "bold"))
        cols_frame.pack(fill=tk.X, padx=10, pady=5)
        labels = [g["label"] for g in groups]
        combo_holder = tk.Frame(cols_frame)
        combo_holder.pack(fill=tk.X)
        col_combos = []

        def add_combo(default_idx=None):
            row = tk.Frame(combo_holder)
            row.pack(fill=tk.X, pady=2)
            tk.Label(row, text=f"Столбец {len(col_combos) + 1}:", width=10, anchor="w").pack(side=tk.LEFT)
            cb = ttk.Combobox(row, state="readonly", width=32, values=labels)
            cb.pack(side=tk.LEFT, padx=4)
            if default_idx is not None and default_idx < len(labels):
                cb.current(default_idx)

            def remove():
                if len(col_combos) <= 2:
                    return
                row.destroy()
                col_combos.remove(cb)

            tk.Button(row, text="−", width=2, command=remove).pack(side=tk.LEFT)
            col_combos.append(cb)

        add_combo(0)
        add_combo(1)
        tk.Button(cols_frame, text="+  Добавить столбец",
                  command=lambda: add_combo(min(len(col_combos), len(labels) - 1))).pack(anchor="w", padx=4, pady=2)
        if len(groups) == 2:
            tk.Label(cols_frame, text="(для Манна-Уитни сравниваются два столбца; "
                                      "для Краскелла-Уоллиса добавьте 3+ через «+»)",
                     fg="#555555").pack(anchor="w", padx=4)

        # Переключатель: показывать значения p с 4 знаками после запятой
        four_var = tk.BooleanVar(value=False)
        tk.Checkbutton(content, text="Показывать p с 4 знаками после запятой",
                       variable=four_var).pack(anchor="w", padx=12, pady=(4, 0))
        always_var = tk.BooleanVar(value=bool(block.get("always_label_criterion", False)))
        tk.Checkbutton(content, text="Всегда подписывать какой критерий использовался",
                       variable=always_var).pack(anchor="w", padx=12, pady=(0, 8))

        def do_calc():
            chosen = [n for n, v in crit_vars.items() if v.get()]
            if not chosen:
                messagebox.showwarning("Критерий", "Выберите критерий для расчёта p")
                return
            criterion = chosen[0]
            sel_feats = [f for f, v in feat_vars.items() if v.get()]
            if not sel_feats:
                messagebox.showwarning("Признаки", "Выберите хотя бы один признак-строку")
                return
            lab2def = {g["label"]: g for g in groups}
            picked = []
            for cb in col_combos:
                gd = lab2def.get(cb.get())
                if gd and gd not in picked:
                    picked.append(gd)
            kw = (criterion == "Критерий Краскелла-Уоллиса")
            if kw:
                if len(picked) < 2:
                    messagebox.showwarning("Столбцы", "Выберите минимум два РАЗНЫХ столбца")
                    return
            else:
                if len(picked) != 2:
                    messagebox.showwarning("Столбцы",
                                           "Для критерия Манна-Уитни нужно ровно два РАЗНЫХ столбца")
                    return
            pv = self._pretty_value
            four = four_var.get()
            small = block.get("small_features", set())
            res_pv, res_cr = {}, {}

            def col_values(feat, gd):
                return df[(df[gd["col"]].astype(str).map(pv) == gd["value"]) &
                          df[feat].notna() &
                          (df[feat].astype(str).str.strip() != '')][feat].astype(float).tolist()

            for feat in sel_feats:
                if feat in small:
                    res_pv[feat] = "-"
                    continue
                if kw:
                    p = self._kruskal_wallis_p([col_values(feat, gd) for gd in picked])
                    res_cr[feat] = "Критерий Краскелла-Уоллиса"
                else:
                    p = self._mann_whitney_p(col_values(feat, picked[0]), col_values(feat, picked[1]))
                    res_cr[feat] = "Критерий Манна-Уитни"
                res_pv[feat] = self._fmt_pvalue(p, four=four)
            for feat in small:
                res_pv.setdefault(feat, "-")
            block.setdefault("p_results", []).append(
                {"pvalues": res_pv, "p_criteria": res_cr,
                 "compared": tuple(gd["value"] for gd in picked)})
            block["always_label_criterion"] = bool(always_var.get())
            self._rerender_all(window)
            dlg.destroy()

        tk.Button(btns, text="Рассчитать p", command=do_calc,
                  bg="green", fg="white").pack(side=tk.LEFT, padx=4)
        tk.Button(btns, text="Отмена", command=dlg.destroy).pack(side=tk.RIGHT, padx=4)

    def _compute_quant_rows(self, df, qual_selected, quant_selected):
        """Возвращает (columns, rows, small_features, group_counts) для таблицы медиан.
        Если N<=4 — вместо медианы перечисляем значения через '; '."""
        qual_unique_values = {}
        for col in qual_selected:
            uv = df[col].dropna().astype(str).unique()
            qual_unique_values[col] = sorted([v for v in uv if str(v).strip() != ''])

        columns = ["Метрика"]
        for col in qual_selected:
            for val in qual_unique_values[col]:
                columns.append(f"N {col}={val}")
                columns.append(f"{col}={val}")

        def fmt_one(v):
            ip, fr = self._dec_parts(v)
            return ip if not fr else ip + ',' + fr

        rows = []
        small_features = set()
        group_counts = {}  # feature -> {group_label: N}
        for quant_col in quant_selected:
            group_counts[quant_col] = {}
            # 1-й проход: считаем статистики по всем группам строки
            cells = []  # None / 'ошибка' / dict (med/q1/q3) / dict (raw)
            for col in qual_selected:
                for val in qual_unique_values[col]:
                    glabel = f"{col} = {val}"
                    try:
                        filtered = df[(df[col].astype(str) == val) &
                                      (df[quant_col].notna()) & (df[quant_col] != '')]
                        n = len(filtered)
                        group_counts[quant_col][glabel] = n
                        if n == 0:
                            cells.append(None)
                        elif n <= 4:
                            vals = sorted(filtered[quant_col].astype(float).tolist())
                            cells.append({"count": n, "raw": vals})
                            small_features.add(quant_col)
                        else:
                            vals = filtered[quant_col].astype(float).to_numpy()
                            cells.append({
                                "count": n,
                                "med": round(float(np.median(vals)), 4),
                                "q1": round(float(np.quantile(vals, 0.25, method='averaged_inverted_cdf')), 4),
                                "q3": round(float(np.quantile(vals, 0.75, method='averaged_inverted_cdf')), 4),
                            })
                    except Exception as e:
                        print(f"Error {quant_col} {col}={val}: {e}")
                        cells.append("ошибка")

            # Ширина строки = макс. число знаков, но число с 5 на конце НЕ учитывается.
            # Сырые значения (N<=4) в расчёте ширины не участвуют.
            parts = []
            for c in cells:
                if isinstance(c, dict) and "med" in c:
                    parts += [self._dec_parts(c["med"]), self._dec_parts(c["q1"]),
                              self._dec_parts(c["q3"])]
            width = max((self._contribution(fr) for ip, fr in parts), default=0)

            # 2-й проход: формируем ячейки
            row_data = [quant_col]
            for c in cells:
                if isinstance(c, dict) and "med" in c:
                    row_data.append(str(c["count"]))
                    me = self._fmt_parts(*self._dec_parts(c["med"]), width)
                    q1 = self._fmt_parts(*self._dec_parts(c["q1"]), width)
                    q3 = self._fmt_parts(*self._dec_parts(c["q3"]), width)
                    row_data.append(f"{me} [{q1}; {q3}]")
                elif isinstance(c, dict) and "raw" in c:
                    row_data.append(str(c["count"]))
                    row_data.append("; ".join(fmt_one(v) for v in c["raw"]))
                elif c == "ошибка":
                    row_data.extend(["ошибка", "ошибка"])
                else:
                    row_data.extend(["0", "-"])
            rows.append(row_data)
        return columns, rows, small_features, group_counts

    def _compute_mean_rows(self, df, qual_selected, quant_selected):
        """Возвращает (columns, rows) для таблицы среднее ± стандартное отклонение (M ± SD).
        Только для количественных признаков-строк."""
        qual_unique_values = {}
        for col in qual_selected:
            uv = df[col].dropna().astype(str).unique()
            qual_unique_values[col] = sorted([v for v in uv if str(v).strip() != ''])

        columns = ["Метрика"]
        for col in qual_selected:
            for val in qual_unique_values[col]:
                columns.append(f"N {col}={val}")
                columns.append(f"{col}={val}")

        rows = []
        for quant_col in quant_selected:
            cells = []
            for col in qual_selected:
                for val in qual_unique_values[col]:
                    try:
                        filtered = df[(df[col].astype(str) == val) &
                                      (df[quant_col].notna()) & (df[quant_col] != '')]
                        n = len(filtered)
                        if n == 0:
                            cells.append(None)
                        else:
                            vals = filtered[quant_col].astype(float).to_numpy()
                            m = float(np.mean(vals))
                            sd = float(np.std(vals, ddof=1)) if n >= 2 else 0.0
                            cells.append({"count": n, "mean": round(m, 4), "sd": round(sd, 4)})
                    except Exception as e:
                        print(f"Error mean {quant_col} {col}={val}: {e}")
                        cells.append("ошибка")
            # выравнивание числа знаков по всем M и SD строки (правило «пятёрки»)
            parts = []
            for c in cells:
                if isinstance(c, dict):
                    parts += [self._dec_parts(c["mean"]), self._dec_parts(c["sd"])]
            width = max((self._contribution(fr) for ip, fr in parts), default=0)

            row_data = [quant_col]
            for c in cells:
                if isinstance(c, dict):
                    row_data.append(str(c["count"]))
                    m = self._fmt_parts(*self._dec_parts(c["mean"]), width)
                    sd = self._fmt_parts(*self._dec_parts(c["sd"]), width)
                    row_data.append(f"{m} ± {sd}")
                elif c == "ошибка":
                    row_data.extend(["ошибка", "ошибка"])
                else:
                    row_data.extend(["0", "-"])
            rows.append(row_data)
        return columns, rows

    def _render_block(self, window, block):
        """Рисует один блок-таблицу в области результатов (со своими столбцами)."""
        frame = tk.LabelFrame(window.results_inner, text=block.get("kind_label", ""),
                              font=("Arial", 9, "bold"), padx=4, pady=4)
        frame.pack(fill=tk.X, expand=False, anchor="w", padx=4, pady=6)

        # Галочка выбора таблицы для объединения (с № выбора)
        sel_var = tk.BooleanVar(value=block.get("_selected", False))

        def on_sel(b=block, v=sel_var, w=window):
            b["_selected"] = v.get()
            order = w.merge_sel_order
            if v.get():
                if b not in order:
                    order.append(b)
            else:
                if b in order:
                    order.remove(b)
            self._refresh_merge_labels(w)

        sel_cb = tk.Checkbutton(frame, text="Выбрать для объединения", variable=sel_var,
                                command=on_sel)
        sel_cb.pack(anchor="w")
        block["_sel_var"] = sel_var
        block["_sel_cb"] = sel_cb

        # Кнопка повторного расчёта p (если столбцов-групп больше двух и таблица не объединённая)
        if not block.get("merged") and not block.get("no_p") and self._block_n_groups(block) > 2:
            tk.Button(frame, text="Рассчитать p повторно",
                      command=lambda b=block, w=window: self._recalc_pvalues(w, b),
                      bg="#1565c0", fg="white").pack(anchor="w", pady=(0, 2))

        specs = [] if block.get("no_p") else self._p_column_specs(block)

        xbar = ttk.Scrollbar(frame, orient=tk.HORIZONTAL)
        xbar.pack(side=tk.BOTTOM, fill=tk.X)

        if block["kind"] == "quant":
            base_cols = list(block["columns"])
            p_ids = [f"p{i}" for i in range(len(specs))]
            columns = base_cols + p_ids
            rows = block["rows"]
            tv = ttk.Treeview(frame, columns=columns, show="headings",
                              height=max(1, len(rows)), xscrollcommand=xbar.set)
            for col in base_cols:
                tv.heading(col, text=col)
                if col.startswith("N "):
                    tv.column(col, width=55, minwidth=45, anchor="center")
                elif col == "Метрика":
                    tv.column(col, width=150, minwidth=100)
                else:
                    tv.column(col, width=150, minwidth=90)
            for pi, sp in enumerate(specs):
                tv.heading(p_ids[pi], text=sp["header"])
                tv.column(p_ids[pi], width=90, minwidth=55, anchor="center")
            for i, r in enumerate(rows):
                feat = r[0] if r else ""
                row_vals = list(r) + [sp["pvalues"].get(feat, "") for sp in specs]
                tv.insert("", tk.END, values=row_vals, tags=('evenrow' if i % 2 == 0 else 'oddrow',))
        else:
            model = block["model"]
            groups = model["groups"]
            col_ids = ["priznak", "kategoria"]
            for gi in range(len(groups)):
                col_ids += [f"N_{gi}", f"np_{gi}"]
            p_ids = [f"p{i}" for i in range(len(specs))]
            col_ids += p_ids
            flat = self._count_flat_rows(model, specs)
            tv = ttk.Treeview(frame, columns=col_ids, show="headings",
                              height=max(1, len(flat)), xscrollcommand=xbar.set)
            tv.heading("priznak", text="Признак"); tv.column("priznak", width=170, minwidth=110)
            tv.heading("kategoria", text="Категория"); tv.column("kategoria", width=110, minwidth=70)
            vh = block.get("value_header", "n (%)")
            for gi, g in enumerate(groups):
                tv.heading(f"N_{gi}", text=f"N ({g['label']})")
                tv.column(f"N_{gi}", width=55, minwidth=45, anchor="center")
                tv.heading(f"np_{gi}", text=f"{vh} ({g['label']})")
                tv.column(f"np_{gi}", width=130, minwidth=80)
            for pi, sp in enumerate(specs):
                tv.heading(p_ids[pi], text=sp["header"])
                tv.column(p_ids[pi], width=90, minwidth=55, anchor="center")
            for i, r in enumerate(flat):
                tv.insert("", tk.END, values=r, tags=('evenrow' if i % 2 == 0 else 'oddrow',))

        tv.tag_configure('evenrow', background='white')
        tv.tag_configure('oddrow', background='#f0f0f0')
        tv.pack(side=tk.TOP, fill=tk.X)
        xbar.config(command=tv.xview)

        if block["kind"] == "count" and block.get("ci"):
            tk.Label(frame, text=block["ci"]["text"], font=("Arial", 9),
                     fg="#222222", anchor="w", justify=tk.LEFT).pack(side=tk.TOP, fill=tk.X, pady=(3, 0))

    def _count_flat_rows(self, model, specs=None):
        """Преобразует модель частот в плоские строки для Treeview.
        p ставится в первой строке признака; столбцов p может быть несколько."""
        specs = specs or [{"pvalues": {}}]
        flat = []
        for f in model["features"]:
            if f.get("no_p"):
                p_txts = ["" for _ in specs]
            else:
                p_txts = [sp["pvalues"].get(f["name"], "") for sp in specs]
            for ri, row in enumerate(f["rows"]):
                values = [f["name"] if ri == 0 else "", row["category"]]
                for cell in row["cells"]:
                    values.append("" if cell["N"] is None else str(cell["N"]))
                    values.append(cell["np"])
                values += [(pt if ri == 0 else "") for pt in p_txts]
                flat.append(values)
        return flat

    def _recalc_pvalues(self, window, block):
        """Открывает то же окно расчёта p для повторного расчёта (добавит новый столбец p)."""
        if block.get("merged"):
            return
        if block["kind"] == "quant":
            self._open_pvalue_dialog(window.df, window, block)
        else:
            self._open_count_pvalue_dialog(window, block)

    @staticmethod
    def _pretty_value(v):
        """'1.0' -> '1'."""
        s = str(v).strip()
        if s.endswith('.0'):
            try:
                f = float(s)
                if f == int(f):
                    return str(int(f))
            except Exception:
                pass
        return s

    @staticmethod
    def _fmt(x, dec=2):
        """Число с десятичной ЗАПЯТОЙ и заданным числом знаков.
        dec=0 -> целое (1), dec=2 -> 1,60 и т.п."""
        return f"{x:.{dec}f}".replace('.', ',')

    @staticmethod
    def _dec_parts(v, cap=4):
        """Разбивает число на целую часть и значащие знаки после запятой
        (после округления максимум до cap знаков, без хвостовых нулей)."""
        s = f"{round(float(v), cap):.{cap}f}".rstrip('0')
        if s.endswith('.'):
            return s[:-1], ''
        ip, fr = s.split('.', 1)
        return ip, fr

    @staticmethod
    def _contribution(frac):
        """Сколько знаков после запятой это число «навязывает» строке.
        Если последний знак 5 — он не учитывается (но само число не меняется)."""
        n = len(frac)
        return n - 1 if (n >= 1 and frac[-1] == '5') else n

    @staticmethod
    def _fmt_parts(ip, frac, width):
        """Формат с запятой: дополняет нулями до width, но не короче собственной длины."""
        d = max(width, len(frac))
        if d == 0:
            return ip
        return ip + ',' + frac + '0' * (d - len(frac))

    _POSITIVE_TOKENS = {'1', 'yes', 'да', 'true', '+', 'positive', 'есть', 'наличие', 'y'}

    def _positive_category(self, cats):
        if len(cats) != 2:
            return None
        pos = [c for c in cats if self._pretty_value(c).lower() in self._POSITIVE_TOKENS]
        return pos[0] if len(pos) == 1 else None

    def _fmt_pct(self, cnt, N, pct_manual=False, pct_decimals=0):
        """Форматирует процент. Авто: целые при N<=100, 1 знак при N>100.
        Ручной режим: округление до pct_decimals знаков."""
        pct = (cnt / N * 100) if N > 0 else 0.0
        d = int(pct_decimals) if pct_manual else (0 if N <= 100 else 1)
        if d <= 0:
            return str(round(pct))
        return f"{pct:.{d}f}".replace('.', ',')

    def _compute_count_structure(self, df, qual_selected, quant_selected, only_group=None,
                                 pct_manual=False, pct_decimals=0):
        pv = self._pretty_value
        groups = []
        if only_group is not None:
            # Только одна выбранная группа (один столбец = значение)
            gcol, gval = only_group
            groups.append({'col': gcol, 'value': gval, 'label': f"{gcol} = {gval}"})
        else:
            for qcol in qual_selected:
                vals = sorted([pv(v) for v in df[qcol].dropna().astype(str).unique()
                               if str(v).strip() != ''])
                for v in vals:
                    groups.append({'col': qcol, 'value': v, 'label': f"{qcol} = {v}"})
        features = []
        for feat in quant_selected:
            cats_all = sorted([pv(v) for v in df[feat].dropna().astype(str).unique()
                               if str(v).strip() != ''])
            # Показываем ВСЕ категории признака (например, и yes, и no), без сворачивания
            shown_cats, binary = cats_all, False
            rows = []
            for ci, cat in enumerate(shown_cats):
                cells = []
                for g in groups:
                    gdf = df[df[g['col']].astype(str).map(pv) == g['value']]
                    fn = gdf[gdf[feat].notna() & (gdf[feat].astype(str).str.strip() != '')]
                    N = len(fn)
                    cnt = int((fn[feat].astype(str).map(pv) == cat).sum())
                    pct_s = self._fmt_pct(cnt, N, pct_manual, pct_decimals)
                    cells.append({'N': (N if ci == 0 else None), 'cnt': cnt,
                                  'np': f"{cnt} ({pct_s}%)"})
                rows.append({'category': cat, 'cells': cells})
            features.append({'name': feat, 'binary': binary, 'rows': rows})
        return {'groups': groups, 'features': features}

    # ---------- 95% ДИ (Клоппера-Пирсона), без внешних зависимостей ----------
    @staticmethod
    def _betai(a, b, x):
        """Регуляризованная неполная бета-функция I_x(a,b)."""
        if x <= 0.0:
            return 0.0
        if x >= 1.0:
            return 1.0
        lnbeta = math.lgamma(a + b) - math.lgamma(a) - math.lgamma(b)
        front = math.exp(lnbeta + a * math.log(x) + b * math.log(1.0 - x))

        def betacf(a, b, x):
            MAXIT, EPS, FPMIN = 300, 3e-12, 1e-300
            qab, qap, qam = a + b, a + 1.0, a - 1.0
            c = 1.0
            d = 1.0 - qab * x / qap
            if abs(d) < FPMIN:
                d = FPMIN
            d = 1.0 / d
            h = d
            for m in range(1, MAXIT + 1):
                m2 = 2 * m
                aa = m * (b - m) * x / ((qam + m2) * (a + m2))
                d = 1.0 + aa * d
                if abs(d) < FPMIN:
                    d = FPMIN
                c = 1.0 + aa / c
                if abs(c) < FPMIN:
                    c = FPMIN
                d = 1.0 / d
                h *= d * c
                aa = -(a + m) * (qab + m) * x / ((a + m2) * (qap + m2))
                d = 1.0 + aa * d
                if abs(d) < FPMIN:
                    d = FPMIN
                c = 1.0 + aa / c
                if abs(c) < FPMIN:
                    c = FPMIN
                d = 1.0 / d
                de = d * c
                h *= de
                if abs(de - 1.0) < EPS:
                    break
            return h

        if x < (a + 1.0) / (a + b + 2.0):
            return front * betacf(a, b, x) / a
        return 1.0 - front * betacf(b, a, 1.0 - x) / b

    @classmethod
    def _inv_betai(cls, p, a, b):
        """Обратная к I_x(a,b) по x (бисекция)."""
        lo, hi = 0.0, 1.0
        for _ in range(200):
            mid = (lo + hi) / 2.0
            if cls._betai(a, b, mid) < p:
                lo = mid
            else:
                hi = mid
        return (lo + hi) / 2.0

    @classmethod
    def _clopper_pearson(cls, k, n, alpha=0.05):
        if n == 0:
            return 0.0, 0.0
        lo = 0.0 if k == 0 else cls._inv_betai(alpha / 2.0, k, n - k + 1)
        hi = 1.0 if k == n else cls._inv_betai(1.0 - alpha / 2.0, k + 1, n - k)
        return lo, hi

    def _compute_ci(self, df, feature, category, group_label):
        """Считает долю выбранной категории и 95% ДИ Клоппера-Пирсона."""
        pv = self._pretty_value
        sub = df
        group_txt = ""
        if group_label and group_label != "(весь набор данных)" and " = " in group_label:
            gcol, gval = group_label.split(" = ", 1)
            gcol = gcol.strip()
            if gcol in df.columns:
                sub = df[df[gcol].astype(str).map(pv) == gval.strip()]
                group_txt = f" ({group_label})"
        fn = sub[sub[feature].notna() & (sub[feature].astype(str).str.strip() != '')]
        n = len(fn)
        k = int((fn[feature].astype(str).map(pv) == category).sum())
        if n == 0:
            text = f"95% ДИ: нет данных по признаку «{feature}»"
            return {"feature": feature, "category": category, "k": 0, "n": 0,
                    "p": 0, "lo": 0, "hi": 0, "text": text}
        lo, hi = self._clopper_pearson(k, n)
        p = k / n
        text = (f"Вероятность «{category}» по признаку «{feature}»{group_txt}: "
                f"{self._fmt(p * 100)}%, 95% ДИ ({self._fmt(lo * 100)}%; {self._fmt(hi * 100)}%)")
        return {"feature": feature, "category": category, "k": k, "n": n,
                "p": p, "lo": lo, "hi": hi, "text": text}

    # ---------- Критерий Манна-Уитни (без внешних зависимостей) ----------
    @staticmethod
    def _norm_sf(z):
        """Хвост стандартного нормального распределения P(Z > z)."""
        return 0.5 * math.erfc(z / math.sqrt(2.0))

    def _mann_whitney_p(self, x, y):
        """Двусторонний p критерия Манна-Уитни (асимптотика: поправки на связи и непрерывность)."""
        x = [float(v) for v in x]
        y = [float(v) for v in y]
        n1, n2 = len(x), len(y)
        if n1 == 0 or n2 == 0:
            return None
        combined = sorted([(v, 0) for v in x] + [(v, 1) for v in y])
        ranks = [0.0] * len(combined)
        tie_terms = 0.0
        i = 0
        while i < len(combined):
            j = i
            while j + 1 < len(combined) and combined[j + 1][0] == combined[i][0]:
                j += 1
            avg = (i + j) / 2.0 + 1.0
            t = j - i + 1
            if t > 1:
                tie_terms += t ** 3 - t
            for kk in range(i, j + 1):
                ranks[kk] = avg
            i = j + 1
        R1 = sum(r for r, (v, g) in zip(ranks, combined) if g == 0)
        U1 = R1 - n1 * (n1 + 1) / 2.0
        mu = n1 * n2 / 2.0
        N = n1 + n2
        sigma2 = (n1 * n2 / 12.0) * ((N + 1) - tie_terms / (N * (N - 1)))
        if sigma2 <= 0:
            return 1.0
        z = (abs(U1 - mu) - 0.5) / math.sqrt(sigma2)  # поправка на непрерывность
        if z < 0:
            z = 0.0
        return min(1.0, 2.0 * self._norm_sf(z))

    def _fmt_pvalue(self, p, four=False):
        """Формат p.
        Без галочки «4 знака»: округление до 3 знаков; если выходит 0,000 -> «<0,001».
        С галочкой: округление до 4 знаков; если выходит 0,0000 -> «<0,0001».
        В режиме 3 знаков «<0,0001» не появляется никогда."""
        if p is None:
            return ""
        if four:
            r4 = round(p, 4)
            if r4 != 0:
                return f"{r4:.4f}".replace('.', ',')
            return "<0,0001"
        r3 = round(p, 3)
        if r3 != 0:
            return f"{r3:.3f}".replace('.', ',')
        return "<0,001"

    @staticmethod
    def _fisher_exact_2col(rows):
        """Точный критерий Фишера для таблицы k×2 (двусторонний, по вероятности таблицы).
        k=2 — ТКФ2; k>2 — критерий Фишера-Фримена-Холтона."""
        k = len(rows)
        if k < 2:
            return None
        R = [r[0] + r[1] for r in rows]
        C0 = sum(r[0] for r in rows)
        C1 = sum(r[1] for r in rows)
        N = C0 + C1
        if N == 0:
            return None
        C = [C0, C1]
        lf = lambda n: math.lgamma(n + 1)
        base = sum(lf(r) for r in R) + sum(lf(c) for c in C) - lf(N)

        def logp(cells):
            return base - sum(lf(x) for x in cells)

        obs = []
        for r in rows:
            obs += [r[0], r[1]]
        p_obs = math.exp(logp(obs))

        # Защита от слишком большого перебора (много категорий/большие N)
        est = 1
        for i in range(k - 1):
            est *= (min(R[i], C0) + 1)
            if est > 4_000_000:
                return None

        suffix = [0] * (k + 1)
        for i in range(k - 1, -1, -1):
            suffix[i] = suffix[i + 1] + R[i]
        EPS = 1e-7
        total = [0.0]

        def rec(i, rem, acc):
            if i == k - 1:
                x = rem
                if 0 <= x <= R[i]:
                    p = math.exp(logp(acc + [x, R[i] - x]))
                    if p <= p_obs * (1 + EPS):
                        total[0] += p
                return
            lo = max(0, rem - suffix[i + 1])
            hi = min(R[i], rem)
            for x in range(lo, hi + 1):
                rec(i + 1, rem - x, acc + [x, R[i] - x])

        rec(0, C0, [])
        return min(1.0, total[0])

    # ---------- χ²-распределение и критерии для >2 групп ----------
    @staticmethod
    def _gammq(a, x):
        """Регуляризованная верхняя неполная гамма Q(a,x) = 1 - P(a,x)."""
        if x <= 0 or a <= 0:
            return 1.0
        gln = math.lgamma(a)
        if x < a + 1.0:
            ap = a; s = 1.0 / a; d = s
            for _ in range(2000):
                ap += 1.0; d *= x / ap; s += d
                if abs(d) < abs(s) * 1e-14:
                    break
            return 1.0 - s * math.exp(-x + a * math.log(x) - gln)
        b = x + 1.0 - a; c = 1e300; d = 1.0 / b; h = d
        for i in range(1, 2000):
            an = -i * (i - a); b += 2.0
            d = an * d + b
            if abs(d) < 1e-300:
                d = 1e-300
            c = b + an / c
            if abs(c) < 1e-300:
                c = 1e-300
            d = 1.0 / d; delta = d * c; h *= delta
            if abs(delta - 1.0) < 1e-14:
                break
        return math.exp(-x + a * math.log(x) - gln) * h

    def _chi2_sf(self, x, df):
        """P(χ²_df > x)."""
        return 1.0 if x <= 0 else self._gammq(df / 2.0, x / 2.0)

    def _kruskal_wallis_p(self, groups):
        """Критерий Краскелла-Уоллиса (>=2 групп), p через χ² с k-1 ст.св."""
        groups = [[float(v) for v in g] for g in groups if len(g) > 0]
        k = len(groups)
        if k < 2:
            return None
        allv = [(v, gi) for gi, g in enumerate(groups) for v in g]
        N = len(allv)
        if N < 2:
            return None
        allv.sort()
        ranks = [0.0] * N
        tie = 0.0
        i = 0
        while i < N:
            j = i
            while j + 1 < N and allv[j + 1][0] == allv[i][0]:
                j += 1
            avg = (i + j) / 2.0 + 1.0
            t = j - i + 1
            if t > 1:
                tie += t ** 3 - t
            for m in range(i, j + 1):
                ranks[m] = avg
            i = j + 1
        Rsum = [0.0] * k
        n = [0] * k
        for (v, gi), r in zip(allv, ranks):
            Rsum[gi] += r
            n[gi] += 1
        if any(c == 0 for c in n):
            return None
        H = 12.0 / (N * (N + 1)) * sum(Rsum[gi] ** 2 / n[gi] for gi in range(k)) - 3 * (N + 1)
        if tie > 0:
            denom = (N ** 3 - N)
            if denom > 0:
                H /= (1 - tie / denom)
        if H <= 0:
            return 1.0
        return self._chi2_sf(H, k - 1)

    def _chi2_test(self, table):
        """Критерий хи-квадрат (Пирсона) для таблицы (строки × столбцы), без поправки."""
        rows = len(table)
        cols = len(table[0]) if rows else 0
        if rows < 2 or cols < 2:
            return None
        rt = [sum(r) for r in table]
        ct = [sum(table[i][j] for i in range(rows)) for j in range(cols)]
        N = sum(rt)
        if N == 0:
            return None
        chi = 0.0
        for i in range(rows):
            for j in range(cols):
                E = rt[i] * ct[j] / N
                if E > 0:
                    chi += (table[i][j] - E) ** 2 / E
        df = (rows - 1) * (cols - 1)
        if df <= 0:
            return None
        return self._chi2_sf(chi, df)

    # ---------- Word ----------
    def _set_word_cell(self, cell, text, bold=False, align=1):
        """Текст в ячейку Word + жирность/выравнивание (0 лево, 1 центр, 2 право)."""
        cell.text = "" if text is None else str(text)
        para = cell.paragraphs[0]
        para.alignment = align
        if cell.text != "":
            for run in para.runs:
                run.font.bold = bold

    def _set_pcell(self, cell, p_txt, index=None, align=1):
        """Ячейка p: значение + (необязательно) верхний индекс с номером критерия."""
        cell.text = "" if p_txt is None else str(p_txt)
        para = cell.paragraphs[0]
        para.alignment = align
        for run in para.runs:
            run.font.size = Pt(10)
        if index and cell.text != "":
            sup = para.add_run(str(index))
            sup.font.superscript = True
            sup.font.size = Pt(10)

    def _append_quant_table(self, doc, columns, rows, p_specs=None, p_index=None,
                            value_label="Me [Q1; Q3]"):
        """Таблица: Признак | группа(N, значение) | ноль или несколько столбцов p.
        value_label — подпись столбца значения (например, «Me [Q1; Q3]» или «M ± SD»)."""
        if p_specs is None:
            p_specs = [{"header": "p", "pvalues": {}, "p_criteria": {}}]
        p_index = p_index or {}
        groups = []
        i = 1
        while i + 1 < len(columns):
            vh = columns[i + 1]
            groups.append(vh.split("=", 1)[1] if "=" in vh else vh)
            i += 2
        n_groups = len(groups)
        n_p = len(p_specs)
        total_cols = 1 + 2 * n_groups + n_p
        n_rows = 2 + len(rows)
        table = doc.add_table(rows=n_rows, cols=total_cols)
        table.style = 'Table Grid'
        table.autofit = False

        priz = table.cell(0, 0).merge(table.cell(1, 0))
        self._set_word_cell(priz, "Признак", True, 1)
        col = 1
        for g in groups:
            gh = table.cell(0, col).merge(table.cell(0, col + 1))
            self._set_word_cell(gh, g, True, 1)
            self._set_word_cell(table.cell(1, col), "N", True, 1)
            self._set_word_cell(table.cell(1, col + 1), value_label, True, 1)
            col += 2
        p_start = 1 + 2 * n_groups
        for pi, sp in enumerate(p_specs):
            pcell = table.cell(0, p_start + pi).merge(table.cell(1, p_start + pi))
            self._set_word_cell(pcell, sp["header"], True, 1)

        for r_idx, values in enumerate(rows):
            rr = 2 + r_idx
            values = list(values)
            feat = values[0] if values else ""
            self._set_word_cell(table.cell(rr, 0), feat, False, 0)
            for k in range(n_groups):
                ni, mi = 1 + 2 * k, 2 + 2 * k
                self._set_word_cell(table.cell(rr, 1 + 2 * k), values[ni] if ni < len(values) else "", False, 1)
                self._set_word_cell(table.cell(rr, 2 + 2 * k), values[mi] if mi < len(values) else "", False, 0)
            for pi, sp in enumerate(p_specs):
                p_txt = sp["pvalues"].get(feat, "")
                idx = p_index.get(sp["p_criteria"].get(feat))
                self._set_pcell(table.cell(rr, p_start + pi), p_txt, idx, 1)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        p_ws = [Inches(1.5) if len(sp["header"]) > 1 else Inches(0.9) for sp in p_specs]
        widths = [Inches(1.7)] + [Inches(0.5), Inches(1.4)] * n_groups + p_ws
        for c_idx in range(total_cols):
            w = widths[c_idx] if c_idx < len(widths) else Inches(1.0)
            for r_idx in range(n_rows):
                try:
                    table.cell(r_idx, c_idx).width = w
                except Exception:
                    pass

    def _append_count_table(self, doc, model, pvalues=None, p_header="p", p_indices=None,
                            value_header="n (%)"):
        """Добавляет в документ таблицу частот (формат примера: блок 'Признак', группы над N/n(%), p)."""
    def _append_count_table(self, doc, model, p_specs=None, p_index=None, value_header="n (%)"):
        """Таблица частот: блок 'Признак', группы над N/value, один или несколько столбцов p."""
        p_specs = p_specs or [{"header": "p", "pvalues": {}, "p_criteria": {}}]
        p_index = p_index or {}
        groups = model['groups']
        features = model['features']
        n_groups = len(groups)
        n_p = len(p_specs)
        total_cols = 2 + 2 * n_groups + n_p
        n_data = sum(len(f['rows']) for f in features)
        n_rows = 2 + n_data
        table = doc.add_table(rows=n_rows, cols=total_cols)
        table.style = 'Table Grid'
        table.autofit = False

        head_priz = table.cell(0, 0).merge(table.cell(1, 1))
        self._set_word_cell(head_priz, "Признак", True, 1)
        p_start = 2 + 2 * n_groups
        for pi, sp in enumerate(p_specs):
            hp = table.cell(0, p_start + pi).merge(table.cell(1, p_start + pi))
            self._set_word_cell(hp, sp["header"], True, 1)
        col = 2
        for g in groups:
            gh = table.cell(0, col).merge(table.cell(0, col + 1))
            self._set_word_cell(gh, g['label'], True, 1)
            self._set_word_cell(table.cell(1, col), "N", True, 1)
            self._set_word_cell(table.cell(1, col + 1), value_header, True, 1)
            col += 2

        r = 2
        for f in features:
            n_sub = len(f['rows'])
            start = r
            for row in f['rows']:
                if not f['binary']:
                    self._set_word_cell(table.cell(r, 1), row['category'], False, 0)
                c = 2
                for cell in row['cells']:
                    self._set_word_cell(table.cell(r, c), "" if cell['N'] is None else cell['N'], False, 1)
                    self._set_word_cell(table.cell(r, c + 1), cell['np'], False, 0)
                    c += 2
                r += 1
            if f['binary']:
                nm = table.cell(start, 0).merge(table.cell(start, 1))
                self._set_word_cell(nm, f['name'], False, 0)
                for pi, sp in enumerate(p_specs):
                    if f.get("no_p"):
                        self._set_pcell(table.cell(start, p_start + pi), "", None, 1)
                    else:
                        self._set_pcell(table.cell(start, p_start + pi), sp["pvalues"].get(f['name'], ""),
                                        p_index.get(sp["p_criteria"].get(f['name'])), 1)
            else:
                nm = table.cell(start, 0)
                for k in range(1, n_sub):
                    nm = nm.merge(table.cell(start + k, 0))
                self._set_word_cell(nm, f['name'], False, 0)
                c = 2
                for gi in range(n_groups):
                    ncell = table.cell(start, c)
                    for k in range(1, n_sub):
                        ncell = ncell.merge(table.cell(start + k, c))
                    nval = f['rows'][0]['cells'][gi]['N']
                    self._set_word_cell(ncell, "" if nval is None else nval, False, 1)
                    c += 2
                for pi, sp in enumerate(p_specs):
                    pcell = table.cell(start, p_start + pi)
                    for k in range(1, n_sub):
                        pcell = pcell.merge(table.cell(start + k, p_start + pi))
                    if f.get("no_p"):
                        self._set_pcell(pcell, "", None, 1)
                    else:
                        self._set_pcell(pcell, sp["pvalues"].get(f['name'], ""),
                                        p_index.get(sp["p_criteria"].get(f['name'])), 1)

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        p_ws = [Inches(1.4) if len(sp["header"]) > 1 else Inches(0.7) for sp in p_specs]
        val_w = Inches(1.3) if len(value_header) > 6 else Inches(1.05)
        widths = [Inches(1.5), Inches(0.9)] + [Inches(0.45), val_w] * n_groups + p_ws
        for c_idx in range(total_cols):
            w = widths[c_idx] if c_idx < len(widths) else Inches(1.0)
            for r_idx in range(n_rows):
                try:
                    table.cell(r_idx, c_idx).width = w
                except Exception:
                    pass

    def export_to_word(self, df, window):
        """Выгружает ВСЕ добавленные таблицы в один документ Word (одна под другой)."""
        blocks = getattr(window, "result_blocks", [])
        if not blocks:
            messagebox.showwarning("Предупреждение",
                                   "Сначала добавьте хотя бы одну таблицу (кнопки «+ Таблица …»)")
            return

        file_path = filedialog.asksaveasfilename(
            title="Сохранить результаты в Word",
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")])
        if not file_path:
            return

        try:
            doc = Document()
            for idx, block in enumerate(blocks):
                heading = doc.add_paragraph()
                label = block.get("kind_label", block.get("title", "Таблица"))
                run = heading.add_run(f"{idx + 1}. {label}")
                run.font.bold = True
                run.font.size = Pt(11)

                # Столбцы p (может быть несколько) и нумерация критериев
                specs = [] if block.get("no_p") else self._p_column_specs(block)
                crit_index, used = self._p_criteria_index(
                    specs, block.get("always_label_criterion", False))

                if block["kind"] == "quant":
                    self._append_quant_table(doc, block["columns"], block["rows"],
                                             specs, crit_index,
                                             block.get("value_header", "Me [Q1; Q3]"))
                else:
                    self._append_count_table(doc, block["model"], specs, crit_index,
                                             block.get("value_header", "n (%)"))
                    if block.get("ci"):
                        ci_p = doc.add_paragraph()
                        ci_run = ci_p.add_run(block["ci"]["text"])
                        ci_run.font.size = Pt(10)

                # Расшифровка индексов критериев под таблицей (если критериев больше одного)
                if crit_index:
                    for c in used:
                        fp = doc.add_paragraph()
                        sup = fp.add_run(str(crit_index[c]))
                        sup.font.superscript = True
                        sup.font.size = Pt(9)
                        nm = fp.add_run(" — " + c)
                        nm.font.size = Pt(9)

                if idx != len(blocks) - 1:
                    doc.add_paragraph("")

            doc.save(file_path)
            messagebox.showinfo("Успех", f"Все таблицы ({len(blocks)}) экспортированы в:\n{file_path}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось экспортировать в Word: {str(e)}")

    def center_window(self):
        self.root.update_idletasks()
        x = (self.root.winfo_screenwidth() // 2) - (900 // 2)
        y = (self.root.winfo_screenheight() // 2) - (600 // 2)
        self.root.geometry(f"900x600+{x}+{y}")
    
    def load_excel_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите Excel файл",
            filetypes=[("Excel files", "*.xlsx *.xls")]
        )
        
        if file_path:
            self.load_button.config(state=tk.DISABLED)
            self.progress.start()
            self.status_label.config(text="Загрузка файла...")
            
            threading.Thread(target=self._load_excel_thread, args=(file_path,), daemon=True).start()
    
    def _load_excel_thread(self, file_path):
        try:
            time.sleep(1)
            self.file_path = file_path
            self.excel_data = pd.read_excel(file_path, sheet_name=None)
            self.root.after(0, self._update_ui_after_load)
        except Exception as e:
            error_msg = str(e)
            self.root.after(0, lambda msg=error_msg: self._show_error(f"Ошибка загрузки файла: {msg}"))
    
    def _update_ui_after_load(self):
        self.progress.stop()
        self.save_excel_button.config(state=tk.NORMAL)
        self.load_button.config(state=tk.NORMAL)
        self.status_label.config(text=f"Загружено листов: {len(self.excel_data)}")
        
        # Очищаем предыдущие кнопки листов
        for widget in self.sheets_frame.winfo_children():
            widget.destroy()
        
        self.sheet_buttons.clear()
        
        # Создаем кнопки для каждого листа
        for sheet_name in self.excel_data.keys():
            btn = tk.Button(self.sheets_frame, text=sheet_name,
                          command=lambda name=sheet_name: self.load_sheet(name))
            btn.pack(side=tk.LEFT, padx=(0, 5))
            self.sheet_buttons[sheet_name] = btn
            
            # Добавляем контекстное меню для удаления листа
            btn.bind("<Button-3>", lambda event, name=sheet_name, button=btn: self.show_sheet_context_menu(event, name, button))
        
        self.save_excel_button.config(state=tk.NORMAL)
    
    def _show_error(self, message):
        self.progress.stop()
        self.load_button.config(state=tk.NORMAL)
        self.status_label.config(text="Ошибка загрузки")
        messagebox.showerror("Ошибка", message)
    
    def load_sheet(self, sheet_name):
        # Сбрасываем выделение предыдущего листа
        if self.current_sheet and self.current_sheet in self.sheet_buttons:
            self.sheet_buttons[self.current_sheet].config(relief=tk.RAISED, bg="#f0f0f0")
        
        self.current_sheet = sheet_name
        df = self.excel_data[sheet_name]
        self.current_columns = list(df.columns)
        self.original_dataframe = df.copy()
        
        # Выделяем текущий лист
        if sheet_name in self.sheet_buttons:
            self.sheet_buttons[sheet_name].config(relief=tk.SUNKEN, bg="lightblue")
        
        self.display_table(df)
        self.config_button.config(state=tk.NORMAL)
        self.reset_filter_button.config(state=tk.DISABLED)
        self.status_label.config(text=f"Отображается лист: {sheet_name}")
        
        # Сбрасываем выделение столбцов при смене листа
        self.selected_columns = []
        if self.criteria_frame:
            self.close_criteria()
    
    def display_table(self, df):
        """Загружает df в редактируемую/фильтруемую таблицу."""
        self.df_full = df.reset_index(drop=True).copy()
        self.col_filters = {}
        self._active_col = None
        self._render_main_table()
        if not self.df_full.empty:
            self.save_as_sheet_button.config(state=tk.NORMAL)
        else:
            self.save_as_sheet_button.config(state=tk.DISABLED)

    @staticmethod
    def _cell_str(v):
        return "" if pd.isna(v) else str(v)

    def _filters_mask(self, exclude_col=None):
        """Булева маска df_full по всем активным фильтрам (можно исключить один столбец)."""
        df = self.df_full
        mask = pd.Series(True, index=df.index)
        for col, allowed in self.col_filters.items():
            if col == exclude_col or col not in df.columns or allowed is None:
                continue
            mask &= df[col].map(lambda v: self._cell_str(v) in allowed)
        return mask

    def _heading_text(self, col):
        mark = " ★" if col in self.col_filters else ""        # фильтр активен
        active = " ◀" if col == self._active_col else ""       # выбранный столбец
        return str(col) + mark + active

    def _refresh_headings(self):
        for col in list(self.table['columns']):
            self.table.heading(col, text=self._heading_text(col))

    def _set_last_action(self, text):
        if hasattr(self, "last_action_label"):
            self.last_action_label.config(text=text)

    def _render_main_table(self, action=None, prev_rows=None):
        """Перерисовывает таблицу с учётом фильтров; iid строки = индекс df.
        action/prev_rows — для сообщения о последнем действии (по числу строк)."""
        self._cancel_cell_edit()
        self._clear_col_highlight()
        if self.df_full is None:
            return
        df = self.df_full
        view = df[self._filters_mask()]
        self.current_displayed_df = view.copy()

        self.table.delete(*self.table.get_children())
        cols = list(df.columns)
        self.table['columns'] = cols
        self.table['show'] = 'tree headings'   # #0 — номера строк
        self.table.heading('#0', text='№')
        self.table.column('#0', width=55, minwidth=40, stretch=False, anchor='center')
        for col in cols:
            self.table.heading(col, text=self._heading_text(col))
            self.table.column(col, width=110, minwidth=50)
        n = 0
        for pos, (idx, row) in enumerate(view.head(3000).iterrows()):
            values = [self._cell_str(v) for v in row]
            tag = 'evenrow' if pos % 2 == 0 else 'oddrow'
            self.table.insert('', tk.END, iid=str(idx), text=str(pos + 1),
                              values=values, tags=(tag,))
            n += 1
        total = len(view)
        self.reset_filter_button.config(state=(tk.NORMAL if self.col_filters else tk.DISABLED))
        if hasattr(self, "rowcount_label"):
            extra = "" if total <= 3000 else f" (показано {n})"
            self.rowcount_label.config(text=f"Строк: {total}{extra}")
        if action and prev_rows is not None:
            after = total
            if action == "filter":
                self._set_last_action(f"Фильтр применён: {prev_rows} → {after}")
            elif action == "reset":
                self._set_last_action(f"Фильтры сброшены: {prev_rows} → {after}")
            elif action == "delete":
                self._set_last_action(f"Удалено {prev_rows - after} строк. {prev_rows} → {after}")
            elif action == "add":
                self._set_last_action(f"Добавлена строка. {prev_rows} → {after}")
        if self._active_col in cols:
            self.table.after(1, self._highlight_active_column)
        self._update_filter_status()

    # ---------- определение столбца/строки по клику ----------
    def _col_at(self, event):
        col_id = self.table.identify_column(event.x)
        if not col_id or col_id == '#0':
            return None
        try:
            i = int(col_id.replace('#', '')) - 1
        except ValueError:
            return None
        cols = list(self.table['columns'])
        return cols[i] if 0 <= i < len(cols) else None

    def on_table_left(self, event):
        # ЛКМ по заголовку — выделить весь столбец
        if self.table.identify_region(event.x, event.y) == "heading":
            col = self._col_at(event)
            if col is not None:
                self._select_column(col)
                return "break"

    def on_table_right(self, event):
        region = self.table.identify_region(event.x, event.y)
        if region == "heading":
            col = self._col_at(event)
            if col is not None:
                self._active_col = col
                self._refresh_headings()
                self._show_column_menu(col, event.x_root, event.y_root)
                return "break"
        elif region == "cell":
            self._show_row_menu(event)
            return "break"

    def on_table_double(self, event):
        region = self.table.identify_region(event.x, event.y)
        if region == "heading":
            col = self._col_at(event)
            if col is not None:
                self._autosize_column(col)
            return "break"
        if region == "cell":
            self._begin_cell_edit(event)
            return "break"

    def _select_column(self, col):
        """Выделяет столбец: помечает активным и подсвечивает только его ячейки."""
        self._active_col = col
        self.current_column_name = col
        try:
            self.show_column_data(col)
        except Exception:
            pass
        self.table.selection_remove(*self.table.selection())  # не выделяем строки целиком
        self._refresh_headings()
        self._highlight_active_column()

    def _clear_col_highlight(self):
        for w in getattr(self, "_colhl", []):
            try:
                w.destroy()
            except Exception:
                pass
        self._colhl = []

    def _highlight_active_column(self):
        """Рисует рамку-подсветку вокруг ячеек активного столбца (только этот столбец)."""
        self._clear_col_highlight()
        col = self._active_col
        if not col:
            return
        try:
            cols = list(self.table['columns'])
        except Exception:
            cols = []
        if col not in cols:
            return
        kids = self.table.get_children()
        if not kids:
            return
        try:
            # первая и последняя ВИДИМЫЕ ячейки столбца
            first = None
            last = None
            for iid in kids:
                bb = self.table.bbox(iid, col)
                if bb:
                    if first is None:
                        first = bb
                    last = bb
            if not first:
                return
            x, y0, w, _ = first
            y1 = last[1] + last[3]
            height = max(0, y1 - y0)
            color = "#1f6fd6"
            specs = [(x, y0, w, 2), (x, y0 + height - 2, w, 2),
                     (x, y0, 2, height), (x + w - 2, y0, 2, height)]
            overlay = []
            for fx, fy, fw, fh in specs:
                fr = tk.Frame(self.table, bg=color)
                fr.place(x=fx, y=fy, width=fw, height=fh)
                overlay.append(fr)
            self._colhl = overlay
        except Exception:
            self._clear_col_highlight()

    def _autosize_column(self, col):
        """Растягивает столбец по самому широкому тексту (учитывая ВСЕ данные столбца)."""
        if self.df_full is None or col not in self.df_full.columns:
            return
        try:
            import tkinter.font as tkfont
            try:
                f = tkfont.Font(font=self.table.cget("font"))
            except Exception:
                f = tkfont.nametofont("TkDefaultFont")
            hf = tkfont.nametofont("TkHeadingFont")
        except Exception:
            f = hf = None
        header = self._heading_text(col)
        series = self.df_full[col]
        if f is not None:
            w = hf.measure(header) if hf is not None else f.measure(header)
            for v in series.astype(object):
                w = max(w, f.measure(self._cell_str(v)))
            w = min(max(w + 26, 50), 1000)
        else:
            longest = max([len(header)] + [len(self._cell_str(v)) for v in series.astype(object)])
            w = min(max(longest * 8 + 26, 50), 1000)
        self.table.column(col, width=int(w), stretch=False)
        self._highlight_active_column()

    # ---------- Фильтры ----------
    def _open_filter_popup(self, col):
        if self.df_full is None or col not in self.df_full.columns:
            return
        # Одновременно открыто только одно окно фильтра
        old = getattr(self, "_filter_popup", None)
        if old is not None:
            try:
                old.destroy()
            except Exception:
                pass
            self._filter_popup = None

        sub = self.df_full[self._filters_mask(exclude_col=col)]
        values = sorted({self._cell_str(v) for v in sub[col]}, key=lambda s: (s == "", s))
        allowed = self.col_filters.get(col)

        pop = tk.Toplevel(self.root)
        self._filter_popup = pop
        pop.title(f"Фильтр: {col}")
        pop.geometry("300x420")
        pop.protocol("WM_DELETE_WINDOW",
                     lambda: (setattr(self, "_filter_popup", None), pop.destroy()))

        topbar = tk.Frame(pop); topbar.pack(fill=tk.X, padx=6, pady=4)
        tk.Button(topbar, text="Выделить всё",
                  command=lambda: [v.set(True) for v in var_map.values()]).pack(side=tk.LEFT)
        tk.Button(topbar, text="Снять всё",
                  command=lambda: [v.set(False) for v in var_map.values()]).pack(side=tk.LEFT, padx=4)

        area = self._scrollable_area(pop)
        var_map = {}
        for val in values:
            checked = (allowed is None) or (val in allowed)
            var = tk.BooleanVar(value=checked)
            var_map[val] = var
            tk.Checkbutton(area, text=("(пусто)" if val == "" else val),
                           variable=var, anchor="w").pack(fill=tk.X, padx=4)

        btns = tk.Frame(pop); btns.pack(side=tk.BOTTOM, fill=tk.X, padx=6, pady=6)

        def close_pop():
            self._filter_popup = None
            pop.destroy()

        def apply_filter():
            chosen = {val for val, var in var_map.items() if var.get()}
            prev = len(self.current_displayed_df) if self.current_displayed_df is not None else 0
            if len(chosen) == len(values):
                self.col_filters.pop(col, None)   # выбрано всё — фильтра нет
            else:
                self.col_filters[col] = chosen
            self._render_main_table(action="filter", prev_rows=prev)
            close_pop()

        def clear_col():
            prev = len(self.current_displayed_df) if self.current_displayed_df is not None else 0
            self.col_filters.pop(col, None)
            self._render_main_table(action="filter", prev_rows=prev)
            close_pop()

        tk.Button(btns, text="Применить", command=apply_filter,
                  bg="green", fg="white").pack(side=tk.LEFT)
        tk.Button(btns, text="Сбросить столбец", command=clear_col).pack(side=tk.LEFT, padx=4)
        tk.Button(btns, text="Отмена", command=close_pop).pack(side=tk.RIGHT)

    def reset_all_filters(self):
        if not self.col_filters:
            return
        prev = len(self.current_displayed_df) if self.current_displayed_df is not None else 0
        self.col_filters = {}
        self._render_main_table(action="reset", prev_rows=prev)

    # ---------- Контекстное меню строки ----------
    def _popup_menu(self, menu, x, y):
        """Показывает контекстное меню так, чтобы оно закрывалось по клику в стороне/Esc."""
        self._dismiss_menu()
        self._ctx_menu = menu
        menu.bind("<FocusOut>", lambda e: self._dismiss_menu())
        menu.bind("<Escape>", lambda e: self._dismiss_menu())
        menu.tk_popup(x, y)   # без немедленного grab_release — иначе на Linux не закрывается

    def _dismiss_menu(self, *args):
        menu = getattr(self, "_ctx_menu", None)
        if menu is not None:
            try:
                menu.unpost()
            except Exception:
                pass
            try:
                menu.grab_release()
            except Exception:
                pass
            self._ctx_menu = None

    def _show_row_menu(self, event):
        iid = self.table.identify_row(event.y)
        if iid and iid not in self.table.selection():
            self.table.selection_set(iid)
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label="Удалить выделенные строки", command=self.delete_selected_rows)
        self._popup_menu(menu, event.x_root, event.y_root)

    def delete_selected_rows(self):
        sel = self.table.selection()
        if not sel or self.df_full is None:
            return
        idxs = [int(i) for i in sel if str(i).isdigit()]
        prev = len(self.current_displayed_df) if self.current_displayed_df is not None else len(self.df_full)
        self.df_full = self.df_full.drop(index=idxs, errors="ignore").reset_index(drop=True)
        self._render_main_table(action="delete", prev_rows=prev)

    # ---------- Контекстное меню столбца ----------
    def _show_column_menu(self, col, x_root, y_root):
        menu = tk.Menu(self.root, tearoff=0)
        menu.add_command(label="Фильтр…", command=lambda: self._open_filter_popup(col))
        menu.add_command(label="Группировка…", command=lambda: self._open_group_dialog(col))
        menu.add_separator()
        menu.add_command(label="Копировать столбец", command=lambda: self.copy_column(col))
        state = tk.NORMAL if self._col_clipboard else tk.DISABLED
        menu.add_command(label="Вставить столбец (после этого)",
                         command=lambda: self.paste_column(col), state=state)
        menu.add_command(label="Вставить в этот столбец (перезаписать)",
                         command=lambda: self.paste_column_over(col), state=state)
        menu.add_command(label="Переименовать столбец", command=lambda: self.rename_column(col))
        menu.add_separator()
        menu.add_command(label="Удалить столбец", command=lambda: self.delete_column(col))
        self._popup_menu(menu, x_root, y_root)

    def copy_column(self, col):
        if self.df_full is not None and col in self.df_full.columns:
            self._col_clipboard = (str(col), self.df_full[col].copy())

    def paste_column(self, after_col):
        if self.df_full is None or not self._col_clipboard:
            return
        name, series = self._col_clipboard
        new_name = name
        k = 1
        while new_name in self.df_full.columns:
            new_name = f"{name}_копия{k}"
            k += 1
        cols = list(self.df_full.columns)
        pos = cols.index(after_col) + 1 if after_col in cols else len(cols)
        vals = series.reset_index(drop=True)
        vals = vals.reindex(range(len(self.df_full))).reset_index(drop=True)
        self.df_full.insert(pos, new_name, vals.values)
        self._render_main_table()

    def paste_column_over(self, col):
        """Перезаписывает значения выбранного столбца данными из буфера (имя столбца сохраняется)."""
        if self.df_full is None or not self._col_clipboard or col not in self.df_full.columns:
            return
        _, series = self._col_clipboard
        vals = series.reset_index(drop=True).reindex(range(len(self.df_full))).reset_index(drop=True)
        self.df_full[col] = vals.values
        self.col_filters.pop(col, None)   # старый фильтр столбца больше не актуален
        self._render_main_table()

    # ---------- Группировка столбца ----------
    def _open_group_dialog(self, col):
        """Окно с кнопкой «+» для задания групп-правил (x<0, x>10 и т.п.)."""
        if self.df_full is None or col not in self.df_full.columns:
            return
        dlg = tk.Toplevel(self.root)
        dlg.title(f"Группировка: {col}")
        dlg.geometry("420x340")

        tk.Label(dlg, text="Правила групп (используйте x как значение ячейки),\n"
                           "например: x<0   |   x>10 and x<20   |   x=='yes'",
                 justify="left").pack(anchor="w", padx=8, pady=(8, 4))

        rows_holder = tk.Frame(dlg)
        # прокручиваемая область для строк-правил
        area = self._scrollable_area(dlg)
        rule_rows = []  # список (frame, entry)

        def add_rule(initial=""):
            row = tk.Frame(area)
            row.pack(fill=tk.X, padx=4, pady=2)
            tk.Label(row, text=f"Группа {len(rule_rows) + 1}:").pack(side=tk.LEFT)
            ent = tk.Entry(row, width=28)
            ent.insert(0, initial)
            ent.pack(side=tk.LEFT, padx=4)

            def remove():
                row.destroy()
                rule_rows[:] = [(f, e) for (f, e) in rule_rows if f is not row]
                for i, (f, e) in enumerate(rule_rows):
                    for ch in f.winfo_children():
                        if isinstance(ch, tk.Label):
                            ch.config(text=f"Группа {i + 1}:")
                            break

            tk.Button(row, text="−", command=remove, width=2).pack(side=tk.LEFT)
            rule_rows.append((row, ent))

        topbar = tk.Frame(dlg)
        topbar.pack(side=tk.TOP, fill=tk.X, padx=8)  # (под областью прокрутки визуально)
        add_btn = tk.Button(dlg, text="+  Добавить группу", command=lambda: add_rule())
        add_btn.pack(anchor="w", padx=8, pady=4)

        btns = tk.Frame(dlg)
        btns.pack(side=tk.BOTTOM, fill=tk.X, padx=8, pady=8)

        def apply_groups():
            rules = [e.get().strip() for _, e in rule_rows if e.get().strip()]
            if not rules:
                messagebox.showwarning("Группировка", "Добавьте хотя бы одно правило")
                return
            self._apply_grouping(col, rules)
            dlg.destroy()

        tk.Button(btns, text="Применить", command=apply_groups,
                  bg="green", fg="white").pack(side=tk.LEFT)
        tk.Button(btns, text="Отмена", command=dlg.destroy).pack(side=tk.RIGHT)

        add_rule()  # одна группа по умолчанию

    @staticmethod
    def _eval_group_rule(rule, raw):
        """Проверяет, удовлетворяет ли значение правилу (x — значение ячейки)."""
        s = "" if pd.isna(raw) else str(raw)
        try:
            x = float(s)
        except (ValueError, TypeError):
            x = s
        try:
            return bool(eval(rule, {"__builtins__": {}}, {"x": x, "abs": abs, "len": len}))
        except Exception:
            return False

    def _apply_grouping(self, col, rules):
        """Создаёт правее столбца col по одному новому столбцу на каждое правило."""
        if self.df_full is None or col not in self.df_full.columns:
            return
        cols = list(self.df_full.columns)
        pos = cols.index(col) + 1
        for offset, rule in enumerate(rules):
            name = f"{col}: {rule}"
            k = 1
            base = name
            while name in self.df_full.columns:
                name = f"{base} ({k})"
                k += 1
            new_vals = [self._cell_str(v) if self._eval_group_rule(rule, v) else ""
                        for v in self.df_full[col]]
            self.df_full.insert(pos + offset, name, new_vals)
        self._render_main_table()

    def delete_column(self, col):
        if self.df_full is None or col not in self.df_full.columns:
            return
        if not messagebox.askyesno("Удаление", f"Удалить столбец «{col}»?"):
            return
        self.df_full = self.df_full.drop(columns=[col])
        self.col_filters.pop(col, None)
        if self._active_col == col:
            self._active_col = None
        self._render_main_table()

    def rename_column(self, col):
        new = tk.simpledialog.askstring("Переименовать", "Новое имя столбца:", initialvalue=str(col))
        if not new or new == col:
            return
        if new in self.df_full.columns:
            messagebox.showwarning("Имя занято", "Столбец с таким именем уже есть")
            return
        self.df_full = self.df_full.rename(columns={col: new})
        if col in self.col_filters:
            self.col_filters[new] = self.col_filters.pop(col)
        self._render_main_table()

    # ---------- Добавление столбца/строки ----------
    def add_column_dialog(self):
        if self.df_full is None:
            messagebox.showwarning("Нет данных", "Сначала загрузите лист")
            return
        name = tk.simpledialog.askstring("Новый столбец", "Имя нового столбца:")
        if not name:
            return
        if name in self.df_full.columns:
            messagebox.showwarning("Имя занято", "Столбец с таким именем уже есть")
            return
        self.df_full[name] = ""
        self._render_main_table()

    def add_row(self):
        if self.df_full is None:
            messagebox.showwarning("Нет данных", "Сначала загрузите лист")
            return
        prev = len(self.current_displayed_df) if self.current_displayed_df is not None else len(self.df_full)
        self.df_full.loc[len(self.df_full)] = ["" for _ in self.df_full.columns]
        self.df_full = self.df_full.reset_index(drop=True)
        self._render_main_table(action="add", prev_rows=prev)
        kids = self.table.get_children()
        if kids:
            self.table.see(kids[-1])

    # ---------- Правка ячейки ----------
    def _cancel_cell_edit(self):
        if self._cell_editor is not None:
            try:
                self._cell_editor.destroy()
            except Exception:
                pass
            self._cell_editor = None

    def _begin_cell_edit(self, event):
        iid = self.table.identify_row(event.y)
        col = self._col_at(event)
        if not iid or col is None:
            return
        bbox = self.table.bbox(iid, col)
        if not bbox:
            return
        x, y, w, h = bbox
        self._cancel_cell_edit()
        cur = ""
        try:
            cur = self.table.set(iid, col)
        except Exception:
            pass
        editor = tk.Entry(self.table)
        editor.insert(0, cur)
        editor.select_range(0, tk.END)
        editor.focus_set()
        editor.place(x=x, y=y, width=w, height=h)
        self._cell_editor = editor

        def commit(_=None):
            new_val = editor.get()
            try:
                idx = int(iid)
                self.df_full.at[idx, col] = new_val
                self.table.set(iid, col, new_val)
            except Exception:
                pass
            self._cancel_cell_edit()

        editor.bind("<Return>", commit)
        editor.bind("<FocusOut>", commit)
        editor.bind("<Escape>", lambda e: self._cancel_cell_edit())

    # ---------- Поиск ----------
    def search_dialog(self):
        if self.df_full is None:
            return
        dlg = tk.Toplevel(self.root)
        dlg.title("Поиск")
        dlg.geometry("360x180")
        tk.Label(dlg, text="Что искать:").pack(anchor="w", padx=8, pady=(8, 0))
        entry = tk.Entry(dlg, width=40)
        entry.pack(fill=tk.X, padx=8)
        entry.focus_set()
        scope = tk.StringVar(value="all")
        sf = tk.LabelFrame(dlg, text="Область")
        sf.pack(fill=tk.X, padx=8, pady=6)
        tk.Radiobutton(sf, text="Вся таблица", variable=scope, value="all").pack(anchor="w")
        tk.Radiobutton(sf, text="Выделенная строка", variable=scope, value="row").pack(anchor="w")
        tk.Radiobutton(sf, text=f"Активный столбец ({self._active_col or '—'})",
                       variable=scope, value="col").pack(anchor="w")

        def do_find():
            self._search_state = None
            self._find_next(entry.get(), scope.get())

        btns = tk.Frame(dlg); btns.pack(side=tk.BOTTOM, fill=tk.X, padx=8, pady=8)
        tk.Button(btns, text="Найти / Далее", command=lambda: self._find_next(entry.get(), scope.get()),
                  bg="green", fg="white").pack(side=tk.LEFT)
        tk.Button(btns, text="Закрыть", command=dlg.destroy).pack(side=tk.RIGHT)
        entry.bind("<Return>", lambda e: self._find_next(entry.get(), scope.get()))

    def _search_targets(self, scope):
        """Список (iid, col) для поиска в заданной области (по видимым строкам)."""
        cols = list(self.table['columns'])
        rows = list(self.table.get_children())
        targets = []
        if scope == "row":
            sel = self.table.selection()
            rows = [r for r in rows if r in sel] or rows
            for r in rows:
                for c in cols:
                    targets.append((r, c))
        elif scope == "col" and self._active_col in cols:
            for r in rows:
                targets.append((r, self._active_col))
        else:
            for r in rows:
                for c in cols:
                    targets.append((r, c))
        return targets

    def _find_next(self, text, scope):
        if not text:
            return
        targets = self._search_targets(scope)
        if not targets:
            return
        st = self._search_state
        start = 0
        if st and st.get("text") == text and st.get("scope") == scope:
            start = st["pos"] + 1
        found = None
        for i in range(start, len(targets) + start):
            iid, col = targets[i % len(targets)]
            val = self.table.set(iid, col)
            if text.lower() in str(val).lower():
                found = (i % len(targets), iid, col)
                break
        if found is None:
            messagebox.showinfo("Поиск", "Совпадений не найдено")
            self._search_state = None
            return
        pos, iid, col = found
        self._search_state = {"text": text, "scope": scope, "pos": pos}
        self.table.selection_set(iid)
        self.table.focus(iid)
        self.table.see(iid)

    # ---------- Замена (regex) ----------
    def replace_dialog(self):
        if self.df_full is None:
            return
        dlg = tk.Toplevel(self.root)
        dlg.title("Замена символов")
        dlg.geometry("380x230")
        tk.Label(dlg, text="Что заменить (регулярное выражение):").pack(anchor="w", padx=8, pady=(8, 0))
        e_from = tk.Entry(dlg, width=44); e_from.pack(fill=tk.X, padx=8)
        tk.Label(dlg, text="На что заменить (можно использовать группы \\1):").pack(anchor="w", padx=8, pady=(6, 0))
        e_to = tk.Entry(dlg, width=44); e_to.pack(fill=tk.X, padx=8)
        scope = tk.StringVar(value="sel")
        sf = tk.LabelFrame(dlg, text="Область замены"); sf.pack(fill=tk.X, padx=8, pady=6)
        tk.Radiobutton(sf, text="Выделенные строки", variable=scope, value="sel").pack(anchor="w")
        tk.Radiobutton(sf, text=f"Активный столбец ({self._active_col or '—'})",
                       variable=scope, value="col").pack(anchor="w")
        tk.Radiobutton(sf, text="Вся таблица", variable=scope, value="all").pack(anchor="w")

        def do_replace():
            self._do_replace(e_from.get(), e_to.get(), scope.get())
            dlg.destroy()

        btns = tk.Frame(dlg); btns.pack(side=tk.BOTTOM, fill=tk.X, padx=8, pady=8)
        tk.Button(btns, text="Заменить", command=do_replace, bg="green", fg="white").pack(side=tk.LEFT)
        tk.Button(btns, text="Отмена", command=dlg.destroy).pack(side=tk.RIGHT)

    def _do_replace(self, pattern, repl, scope):
        if pattern == "":
            return
        try:
            rx = re.compile(pattern)
        except re.error as e:
            messagebox.showerror("Регулярное выражение", f"Ошибка в шаблоне: {e}")
            return
        cols = list(self.df_full.columns)
        # какие столбцы и строки затрагиваем
        if scope == "col" and self._active_col in cols:
            target_cols = [self._active_col]
            idxs = [int(i) for i in self.table.get_children()]
        elif scope == "sel":
            sel = self.table.selection()
            idxs = [int(i) for i in sel] if sel else [int(i) for i in self.table.get_children()]
            target_cols = [self._active_col] if self._active_col in cols else cols
        else:
            target_cols = cols
            idxs = [int(i) for i in self.table.get_children()]
        count = 0
        for idx in idxs:
            for col in target_cols:
                try:
                    old = self._cell_str(self.df_full.at[idx, col])
                    new = rx.sub(repl, old)
                    if new != old:
                        self.df_full.at[idx, col] = new
                        count += 1
                except Exception:
                    pass
        self._render_main_table()
        messagebox.showinfo("Замена", f"Заменено ячеек: {count}")
    
    def configure_table(self):
        if not self.current_sheet:
            return
        
        dialog = ColumnConfigDialog(self.root, self.current_columns)
        self.root.wait_window(dialog.dialog)
        
        if dialog.result is not None:
            self.current_columns = dialog.result
            df = self.excel_data[self.current_sheet]
            
            if self.current_columns:
                filtered_df = df[self.current_columns]
                self.display_table(filtered_df)
            else:
                messagebox.showwarning("Предупреждение", "Не выбрано ни одного столбца!")
    
    def on_column_click(self, event):
        region = self.table.identify_region(event.x, event.y)
        
        if region == "heading":
            column = self.table.identify_column(event.x)
            if column:
                col_index = int(column.replace('#', '')) - 1
                columns = self.table['columns']
                
                if 0 <= col_index < len(columns):
                    column_name = columns[col_index]
                    self.show_column_data(column_name)
    
    def show_column_data(self, column_name):
        if not self.current_sheet:
            return
        
        self.current_column_name = column_name
        df = self.current_displayed_df if self.current_displayed_df is not None else self.excel_data[self.current_sheet]
        
        if column_name in df.columns:
            # Очищаем таблицу
            self.column_data_table.delete(*self.column_data_table.get_children())
            
            # Получаем данные столбца БЕЗ удаления NaN значений
            column_data = df[column_name]
            
            # Переменная для отслеживания максимальной длины текста
            max_text_length = 0
            
            # Добавляем все значения, включая пустые
            for index, value in enumerate(column_data.head(500)):  # Ограничиваем для производительности
                row_num = index + 1
                
                # Обрабатываем пустые значения
                if pd.isna(value) or value == '' or str(value).strip() == '':
                    display_value = ""  # Пустая строка для пустых значений
                    tag = 'empty_row'
                else:
                    display_value = str(value)
                    tag = 'evenrow' if index % 2 == 0 else 'oddrow'
                    # Обновляем максимальную длину
                    max_text_length = max(max_text_length, len(display_value))
                
                # Вставляем строку в таблицу
                self.column_data_table.insert('', tk.END, 
                                            values=(row_num, display_value), 
                                            tags=(tag,))
            
            # Автоматически подстраиваем ширину столбца под максимальную длину текста
            # Используем приблизительное соотношение: 1 символ ≈ 8 пикселей
            # Минимальная ширина 100 пикселей, максимальная 800 пикселей
            char_width = 8
            min_width = 100
            max_width = 800
            calculated_width = max(min_width, min(max_width, max_text_length * char_width + 20))
            
            self.column_data_table.column('value', width=calculated_width)
            
            # Обновляем размер правой панели под новую ширину столбца
            right_frame_width = max(200, calculated_width + 80)  # 80 пикселей на номер строки + отступы
            # Находим правую панель и обновляем её ширину
            for widget in self.root.winfo_children()[0].winfo_children():
                if isinstance(widget, tk.Frame):
                    for child in widget.winfo_children():
                        if isinstance(child, tk.Frame) and hasattr(child, 'winfo_reqwidth'):
                            if child.winfo_reqwidth() >= 200:  # Это правая панель
                                child.config(width=right_frame_width)
                                break
    
    def clear_data(self):
        """Очищает все данные и сбрасывает интерфейс"""
        if messagebox.askyesno("Подтверждение", "Очистить все данные и закрыть файл?"):
            # Очищаем данные
            self.excel_data.clear()
            self.current_sheet = None
            self.current_columns = []
            
            # Очищаем таблицу
            self.table.delete(*self.table.get_children())
            self.table['columns'] = ()
            
            # Очищаем данные столбца
            self.column_data_table.delete(*self.column_data_table.get_children())
            
            # Очищаем кнопки листов
            for widget in self.sheets_frame.winfo_children():
                widget.destroy()
            
            # Отключаем кнопки
            self.config_button.config(state=tk.DISABLED)
            self.reset_filter_button.config(state=tk.DISABLED)
            self.save_as_sheet_button.config(state=tk.DISABLED)
            self.save_excel_button.config(state=tk.DISABLED)
            
            # Сбрасываем дополнительные переменные
            self.current_column_name = None
            self.original_dataframe = None
            self.current_displayed_df = None
            self.file_path = None
            
            # Обновляем статус
            self.status_label.config(text="Данные очищены. Готов к загрузке нового файла")
    
    def copy_selected_data(self):
        """Копирует выделенные данные в буфер обмена"""
        try:
            selected_items = self.column_data_table.selection()
            if selected_items:
                selected_data = []
                for item in selected_items:
                    values = self.column_data_table.item(item, 'values')
                    if values and len(values) > 1:
                        selected_data.append(values[1])  # Берем значение, а не номер строки
                
                # Копируем в буфер обмена
                text_to_copy = '\n'.join(selected_data)
                self.root.clipboard_clear()
                self.root.clipboard_append(text_to_copy)
                
                messagebox.showinfo("Информация", f"Скопировано {len(selected_data)} элементов")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось скопировать данные: {str(e)}")
    
    def hide_context_menus(self, event):
        """Скрывает все контекстные меню"""
        self._dismiss_menu()
        try:
            self.context_menu.unpost()
        except Exception:
            pass

    def show_context_menu(self, event):
        """Показывает контекстное меню для копирования"""
        if self.column_data_table.selection():
            try:
                self.context_menu.delete(0, tk.END)  # Очищаем предыдущие пункты меню
                self.context_menu.add_command(label="Копировать выделенное", command=self.copy_selected_data)
                self.context_menu.add_command(label="Отобразить выделенное", command=self.filter_table_by_selected_rows)
                self.context_menu.add_command(label="Показать подобные", command=self.filter_table_by_similar_values)
                self.context_menu.tk_popup(event.x_root, event.y_root)
                # Блокируем событие, чтобы меню не закрылось сразу
                self.context_menu.grab_set()
            finally:
                # Освобождаем захват после закрытия меню
                self.context_menu.after(100, self.context_menu.grab_release)
    
    def filter_table_by_similar_values(self):
        """Фильтрует таблицу по вхождению текста с сохранением в стеке фильтров"""
        try:
            selected_items = self.column_data_table.selection()
            if not selected_items or not self.current_column_name or not self.current_sheet:
                messagebox.showwarning("Предупреждение", "Выберите значения для поиска")
                return
            
            # Получаем выделенные значения
            search_values = []
            for item in selected_items:
                values = self.column_data_table.item(item, 'values')
                if values and len(values) > 1:
                    value = values[1]
                    if value:
                        search_values.append(str(value).lower())
            
            if not search_values:
                messagebox.showwarning("Предупреждение", "Не выбрано значений для поиска")
                return
            
            # Берем последний отфильтрованный DataFrame или оригинальный, если стек пуст
            current_df = self.filter_stack[-1][1] if self.filter_stack else self.original_dataframe
            
            # Применяем новый фильтр
            mask = current_df[self.current_column_name].astype(str).str.lower().str.contains('|'.join(search_values), na=False)
            filtered_df = current_df[mask]
            
            if filtered_df.empty:
                messagebox.showwarning("Предупреждение", "Нет строк, содержащих выделенные значения")
                return
            
            # Сохраняем фильтр в стек
            filter_description = f"Поиск в '{self.current_column_name}': {', '.join(search_values[:3])}{'...' if len(search_values) > 3 else ''}"
            self.filter_stack.append((filter_description, filtered_df))
            
            # Применяем фильтр столбцов, если он был настроен
            if self.current_columns and set(self.current_columns).issubset(set(filtered_df.columns)):
                filtered_df = filtered_df[self.current_columns]
            
            # Отображаем отфильтрованную таблицу
            self.display_table(filtered_df)
            
            # Активируем кнопку сброса фильтра
            self.reset_filter_button.config(state=tk.NORMAL)
            
            # Обновляем статус
            self._update_filter_status()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось выполнить поиск: {str(e)}")

    def filter_table_by_selected_rows(self):
        """Фильтрует таблицу по выделенным строкам с сохранением в стеке фильтров"""
        try:
            selected_items = self.column_data_table.selection()
            if not selected_items or not self.current_column_name or not self.current_sheet:
                messagebox.showwarning("Предупреждение", "Выберите строки для фильтрации")
                return
            
            # Получаем номера выделенных строк
            selected_rows = []
            for item in selected_items:
                values = self.column_data_table.item(item, 'values')
                if values and len(values) > 0:
                    row_num = int(values[0]) - 1
                    selected_rows.append(row_num)
            
            # Берем последний отфильтрованный DataFrame или оригинальный, если стек пуст
            current_df = self.filter_stack[-1][1] if self.filter_stack else self.original_dataframe
            
            # Применяем новый фильтр
            filtered_df = current_df.iloc[selected_rows]
            
            if filtered_df.empty:
                messagebox.showwarning("Предупреждение", "Нет данных для отображения")
                return
            
            # Сохраняем фильтр в стек
            filter_description = f"Выделено строк: {len(selected_rows)}"
            self.filter_stack.append((filter_description, filtered_df))
            
            # Применяем фильтр столбцов, если он был настроен
            if self.current_columns and set(self.current_columns).issubset(set(filtered_df.columns)):
                filtered_df = filtered_df[self.current_columns]
            
            # Отображаем отфильтрованную таблицу
            self.display_table(filtered_df)
            
            # Активируем кнопку сброса фильтра
            self.reset_filter_button.config(state=tk.NORMAL)
            
            # Обновляем статус
            self._update_filter_status()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось отобразить выделенные строки: {str(e)}")

    def filter_table_by_selection(self):
        """Отобразить в таблице: фильтрует по выделенным значениям столбца (через общий фильтр)."""
        try:
            selected_items = self.column_data_table.selection()
            if not selected_items or not self.current_column_name or self.df_full is None:
                messagebox.showwarning("Предупреждение", "Выберите значения для фильтрации")
                return
            selected_values = set()
            for item in selected_items:
                values = self.column_data_table.item(item, 'values')
                if values and len(values) > 1:
                    selected_values.add(str(values[1]))
            if not selected_values:
                return
            self.col_filters[self.current_column_name] = selected_values
            self._active_col = self.current_column_name
            prev = len(self.current_displayed_df) if self.current_displayed_df is not None else 0
            self._render_main_table(action="filter", prev_rows=prev)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось применить фильтр: {str(e)}")

    def _update_filter_status(self):
        """Обновляет статусную строку с информацией о применённых фильтрах."""
        try:
            sheet = self.current_sheet or ""
            n = len(self.current_displayed_df) if self.current_displayed_df is not None else 0
            if self.col_filters:
                cols = ", ".join(self.col_filters.keys())
                self.status_label.config(
                    text=f"Лист: {sheet} | фильтры по: {cols} | строк: {n}")
            else:
                self.status_label.config(text=f"Лист: {sheet} (без фильтров) | строк: {n}")
        except Exception:
            pass

    
    def reset_table_filter(self):
        """Сбрасывает все фильтры и показывает исходные данные"""
        if not self.current_sheet:
            return
        
        try:
            # Очищаем стек фильтров
            self.filter_stack = []
            
            # Восстанавливаем оригинальные данные
            df = self.original_dataframe
            
            # Применяем настройки столбцов, если они были
            if self.current_columns and set(self.current_columns).issubset(set(df.columns)):
                df = df[self.current_columns]
            
            # Отображаем полную таблицу
            self.display_table(df)
            
            # Отключаем кнопку сброса фильтра
            self.reset_filter_button.config(state=tk.DISABLED)
            
            # Обновляем статус
            self._update_filter_status()
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сбросить фильтр: {str(e)}")
    
    def save_current_table_as_sheet(self):
        """Сохраняет текущую отображаемую таблицу как новый лист"""
        if self.current_displayed_df is None or self.current_displayed_df.empty:
            messagebox.showwarning("Предупреждение", "Нет данных для сохранения")
            return
        
        # Диалог для ввода имени листа
        sheet_name = tk.simpledialog.askstring("Новый лист", 
                                              "Введите имя для нового листа:",
                                              initialvalue=f"Фильтр_{self.current_sheet}")
        
        if not sheet_name:
            return
        
        # Проверяем, не существует ли уже такой лист
        if sheet_name in self.excel_data:
            if not messagebox.askyesno("Подтверждение", 
                                     f"Лист '{sheet_name}' уже существует. Заменить?"):
                return
        
        try:
            # Добавляем новый лист в данные
            self.excel_data[sheet_name] = self.current_displayed_df.copy()
            
            # Обновляем кнопки листов
            self._refresh_sheet_buttons()
            
            messagebox.showinfo("Успех", f"Лист '{sheet_name}' создан с {len(self.current_displayed_df)} строками")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать лист: {str(e)}")
    
    def show_sheet_context_menu(self, event, sheet_name, button):
        """Показывает контекстное меню для листа"""
        if len(self.excel_data) <= 1:
            messagebox.showwarning("Предупреждение", "Нельзя удалить последний лист")
            return
        
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label=f"Удалить лист '{sheet_name}'", 
                            command=lambda: self.delete_sheet(sheet_name, button))
        
        try:
            context_menu.tk_popup(event.x_root, event.y_root)
            context_menu.grab_set()
        finally:
            context_menu.after(100, context_menu.grab_release)
    
    def delete_sheet(self, sheet_name, button):
        """Удаляет лист"""
        if messagebox.askyesno("Подтверждение", f"Удалить лист '{sheet_name}'?"):
            try:
                # Удаляем из данных
                del self.excel_data[sheet_name]
                
                # Если это был текущий лист, очищаем таблицу
                if self.current_sheet == sheet_name:
                    self.table.delete(*self.table.get_children())
                    self.table['columns'] = ()
                    self.column_data_table.delete(*self.column_data_table.get_children())
                    self.current_sheet = None
                    self.config_button.config(state=tk.DISABLED)
                    self.reset_filter_button.config(state=tk.DISABLED)
                    self.save_as_sheet_button.config(state=tk.DISABLED)
                
                # Удаляем кнопку
                button.destroy()
                
                # Если листов не осталось, отключаем кнопку сохранения Excel
                if not self.excel_data:
                    self.save_excel_button.config(state=tk.DISABLED)
                
                messagebox.showinfo("Успех", f"Лист '{sheet_name}' удален")
                
            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось удалить лист: {str(e)}")
    
    def _refresh_sheet_buttons(self):
        """Обновляет кнопки листов"""
        # Очищаем старые кнопки
        for widget in self.sheets_frame.winfo_children():
            widget.destroy()
        
        self.sheet_buttons.clear()
        
        # Создаем новые кнопки
        for sheet_name in self.excel_data.keys():
            btn = tk.Button(self.sheets_frame, text=sheet_name,
                        command=lambda name=sheet_name: self.load_sheet(name))
            btn.pack(side=tk.LEFT, padx=(0, 5))
            self.sheet_buttons[sheet_name] = btn
            
            # Добавляем контекстное меню для удаления листа
            btn.bind("<Button-3>", lambda event, name=sheet_name, button=btn: self.show_sheet_context_menu(event, name, button))
    
    def save_excel_file(self):
        """Сохраняет Excel файл"""
        if not self.excel_data:
            messagebox.showwarning("Предупреждение", "Нет данных для сохранения")
            return
        
        # Выбираем файл для сохранения
        file_path = filedialog.asksaveasfilename(
            title="Сохранить Excel файл",
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("Excel files", "*.xls")]
        )
        
        if not file_path:
            return
        
        try:
            # Сохраняем все листы в Excel файл
            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                for sheet_name, df in self.excel_data.items():
                    df.to_excel(writer, sheet_name=sheet_name, index=False)
            
            messagebox.showinfo("Успех", f"Файл сохранен: {file_path}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить файл: {str(e)}")
    
    # ================== Автообновление с GitHub ==================
    def _fetch_remote_info(self):
        """Скачивает raw-файл из репозитория, читает его __version__ и размер.
        Возвращает dict {version, size, content(bytes)} либо бросает исключение
        (нет интернета, репозиторий не настроен и т.п.)."""
        if GITHUB_OWNER.startswith("ВАШ_") or GITHUB_REPO.startswith("ВАШ_"):
            raise RuntimeError("Не заданы GITHUB_OWNER / GITHUB_REPO")
        url = _raw_github_url()
        req = urllib.request.Request(url, headers={"User-Agent": "ExcelViewerApp-Updater"})
        with urllib.request.urlopen(req, timeout=UPDATE_TIMEOUT) as resp:
            content = resp.read()
        text = content.decode("utf-8", errors="replace")
        m = re.search(r'__version__\s*=\s*["\']([^"\']+)["\']', text)
        if not m:
            raise RuntimeError("В удалённом файле не найдена строка __version__")
        return {"version": m.group(1), "size": len(content), "content": content}

    def check_updates_on_startup(self):
        """Тихая фоновая проверка при запуске: если версия новее — предложить
        обновиться; если нет интернета или версия последняя — ничего не делать."""
        def worker():
            try:
                info = self._fetch_remote_info()
            except Exception:
                return  # нет интернета / не настроено — молча выходим
            if _version_is_newer(info["version"], __version__):
                self.root.after(0, lambda: self._prompt_update(info))
        threading.Thread(target=worker, daemon=True).start()

    def check_updates_manual(self):
        """Принудительная проверка по кнопке «Проверить обновления»."""
        self.check_updates_button.config(state=tk.DISABLED, text="Проверка...")

        def worker():
            error = None
            info = None
            try:
                info = self._fetch_remote_info()
            except Exception as e:
                error = str(e)

            def done():
                self.check_updates_button.config(state=tk.NORMAL, text="Проверить обновления")
                if error is not None:
                    messagebox.showwarning(
                        "Проверка обновлений",
                        "Не удалось проверить обновления.\n"
                        "Проверьте подключение к интернету и настройки репозитория.\n\n"
                        f"Причина: {error}")
                    return
                if _version_is_newer(info["version"], __version__):
                    self._prompt_update(info)
                else:
                    messagebox.showinfo(
                        "Проверка обновлений",
                        f"У вас установлена последняя версия ({__version__}).")

            self.root.after(0, done)

        threading.Thread(target=worker, daemon=True).start()

    def _prompt_update(self, info):
        """Диалог с запросом на обновление и примерным размером."""
        size_kb = info["size"] / 1024.0
        size_txt = f"{size_kb:.0f} КБ" if size_kb < 1024 else f"{size_kb / 1024:.1f} МБ"
        if messagebox.askyesno(
                "Доступно обновление",
                f"Доступна новая версия: {info['version']}\n"
                f"Текущая версия: {__version__}\n"
                f"Примерный размер обновления: ~{size_txt}\n\n"
                "Обновить сейчас?"):
            self._apply_update(info)

    def _apply_update(self, info):
        """Заменяет текущий .py-файл содержимым из репозитория и предлагает
        перезапустить приложение."""
        try:
            target = os.path.abspath(__file__)
        except NameError:
            target = os.path.abspath(sys.argv[0])
        try:
            # Резервная копия на случай сбоя
            backup = target + ".bak"
            try:
                with open(target, "rb") as f_old:
                    old_bytes = f_old.read()
                with open(backup, "wb") as f_bak:
                    f_bak.write(old_bytes)
            except Exception:
                pass  # бэкап не критичен

            with open(target, "wb") as f_new:
                f_new.write(info["content"])
        except Exception as e:
            messagebox.showerror(
                "Ошибка обновления",
                f"Не удалось записать обновление в файл:\n{target}\n\n{e}")
            return

        if messagebox.askyesno(
                "Обновление установлено",
                f"Обновление до версии {info['version']} установлено.\n"
                "Перезапустить приложение сейчас?"):
            try:
                python = sys.executable
                os.execl(python, python, target, *sys.argv[1:])
            except Exception:
                messagebox.showinfo(
                    "Перезапуск",
                    "Не удалось перезапустить автоматически.\n"
                    "Закройте и запустите приложение заново.")
    # =============================================================

    def run(self):
        # Тихая проверка обновлений при запуске (не мешает, если нет интернета)
        self.root.after(800, self.check_updates_on_startup)
        self.root.mainloop()

if __name__ == "__main__":
    app = ExcelViewerApp()
    app.run()
